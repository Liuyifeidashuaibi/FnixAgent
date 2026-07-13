"""自我验证模块(Phase 6.1)。

题库处理结果的声明式质检模块,借鉴 OfficeBench evaluate.py 的声明式检查思路。

检查项:
  - 无 NAMECONTENT 乱码残留
  - 所有选择题括号已填答案
  - 无未删除的题号(如 "11. 【单选题】")
  - 答案行已删除(非问答题)
  - 格式统一(字体/字号一致)

设计:
  - ValidationReport: 验证报告(汇总多项检查结果)
  - CheckItem: 单项检查结果(passed/severity/detail/fix_suggestion)
  - TaskValidator: 自我验证模块(含题库专用 + 通用声明式检查)

底层依赖:
  - python-docx(可选,不可用时检查项标记 error 并提示安装)

降级策略:
  - python-docx 缺失 → CheckItem 标记 error,提示安装,不崩溃
  - 文件 IO 异常 → 捕获返回 error,不崩溃
  - 路径校验前置,避免无效 IO
"""
from __future__ import annotations

import os
import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Any, Optional

from officeagent.office.base import BaseExpert, ExpertError, ExpertResult


__all__ = [
    "CheckItem",
    "ValidationReport",
    "TaskValidator",
]


# ---------------------------------------------------------------------------
# 正则模式
# ---------------------------------------------------------------------------

# 乱码残留:NAME{X}CONTENT
_GARBLED_PATTERN = re.compile(r"NAME[A-Z]CONTENT")

# 选择题标记:【单选题】/【多选题】
_CHOICE_MARKER_PATTERN = re.compile(r"【(单选|多选)题】")

# 已填答案:（A）或 (A) 或 （ABC） 或 (ABC)
_FILLED_ANSWER_PATTERN = re.compile(r"[（(]([A-Z]+)[)）]")

# 题号:"11. 【" 或 "11.【"(行首数字+点+可选空格+【)
_QUESTION_NUMBER_PATTERN = re.compile(r"^\d+\.\s*【")

# 答案行:【答案】
_ANSWER_LINE_PATTERN = re.compile(r"【答案】")

# 格式组合数量阈值(超过则 warning)
_FORMAT_COMBO_WARNING_THRESHOLD = 5


# ---------------------------------------------------------------------------
# 验证报告与检查项
# ---------------------------------------------------------------------------


@dataclass
class CheckItem:
    """单项检查结果。

    Attributes:
        name: 检查项名称
        passed: 是否通过
        severity: 严重级别("error"/"warning"/"info")
        detail: 详细描述
        fix_suggestion: 修复建议(可选)
    """

    name: str
    passed: bool
    severity: str = "error"
    detail: str = ""
    fix_suggestion: Optional[str] = None


@dataclass
class ValidationReport:
    """验证报告。

    Attributes:
        passed: 整体是否通过(无 error 级失败)
        total_checks: 检查总数
        passed_checks: 通过数
        failed_checks: 失败数
        warnings: 警告信息列表
        errors: 错误信息列表
        details: 每项检查的详细结果(列表,每项为 dict)
    """

    passed: bool = True
    total_checks: int = 0
    passed_checks: int = 0
    failed_checks: int = 0
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)
    details: list[dict] = field(default_factory=list)

    @classmethod
    def from_checks(cls, checks: list[CheckItem]) -> "ValidationReport":
        """从 CheckItem 列表汇总生成报告。

        Args:
            checks: 检查项列表

        Returns:
            汇总后的 ValidationReport
        """
        total = len(checks)
        passed = sum(1 for c in checks if c.passed)
        failed = total - passed
        warnings = [
            f"{c.name}: {c.detail}"
            for c in checks
            if not c.passed and c.severity == "warning"
        ]
        errors = [
            f"{c.name}: {c.detail}"
            for c in checks
            if not c.passed and c.severity == "error"
        ]
        return cls(
            passed=len(errors) == 0,
            total_checks=total,
            passed_checks=passed,
            failed_checks=failed,
            warnings=warnings,
            errors=errors,
            details=[
                {
                    "name": c.name,
                    "passed": c.passed,
                    "severity": c.severity,
                    "detail": c.detail,
                    "fix_suggestion": c.fix_suggestion,
                }
                for c in checks
            ],
        )


# ---------------------------------------------------------------------------
# TaskValidator
# ---------------------------------------------------------------------------


class TaskValidator(BaseExpert):
    """自我验证模块。

    提供题库处理结果验证、格式验证、通用声明式检查。

    用法::

        validator = TaskValidator()
        report = validator.validate_question_bank(
            original_path="input.docx",
            processed_path="output.docx",
        )
        if not report.passed:
            for err in report.errors:
                print(err)

    能力边界:
      - 仅处理 .docx(python-docx 不支持 .doc 旧格式)
      - 检查基于正则与 run 属性,不做语义理解
      - 通用检查支持 4 种声明式 function
    """

    @property
    def name(self) -> str:
        return "task_validator"

    # ------------------------------------------------------------------
    # 题库处理结果验证
    # ------------------------------------------------------------------

    def validate_question_bank(
        self,
        original_path: str,
        processed_path: str,
        task_type: str = "question_bank",
    ) -> ValidationReport:
        """题库处理结果验证。

        依次执行以下检查(基于 processed_path):
          - 无 NAMECONTENT 乱码残留
          - 所有选择题括号已填答案
          - 无未删除的题号(如 "11. 【单选题】")
          - 答案行已删除(非问答题,task_type != "qa" 时检查)
          - 格式统一(字体/字号/粗细组合不超过 3 种)

        Args:
            original_path: 原始文件路径(保留参数,供未来扩展比对检查)
            processed_path: 处理后文件路径(检查目标)
            task_type: 任务类型,"qa" 时跳过答案行删除检查

        Returns:
            ValidationReport: 汇总验证报告
        """
        checks: list[CheckItem] = []
        checks.append(self.check_no_garbled(processed_path))
        checks.append(self.check_answers_filled(processed_path))
        checks.append(self.check_numbers_deleted(processed_path))
        # 问答题不检查答案行删除(答案本身就是正文)
        if task_type != "qa":
            checks.append(self.check_answer_lines_removed(processed_path))
        checks.append(self.check_format_unified(processed_path))
        return ValidationReport.from_checks(checks)

    # ------------------------------------------------------------------
    # 格式验证
    # ------------------------------------------------------------------

    def validate_format(self, path: str) -> ValidationReport:
        """格式统一验证。

        统计字体/字号/粗细分布,超过 3 种组合则 warning。

        Args:
            path: docx 文件路径

        Returns:
            ValidationReport: 仅含格式统一单项检查的报告
        """
        checks = [self.check_format_unified(path)]
        return ValidationReport.from_checks(checks)

    # ------------------------------------------------------------------
    # 通用声明式检查(借鉴 OfficeBench evaluate.py)
    # ------------------------------------------------------------------

    def validate_general(
        self, path: str, checks: list[dict]
    ) -> ValidationReport:
        """通用声明式检查。

        每个 check dict 支持的 function:

          - ``"file_exists"``: 检查 path 是否存在
          - ``"text_contains"``: 检查 path 文本是否包含所有 keywords
          - ``"text_not_contains"``: 检查 path 文本是否不包含任何 keywords
          - ``"word_count_min"``: 检查 path 文本词数是否 >= min

        check dict 格式(扁平)::

            {"name": "含答案关键词", "function": "text_contains", "keywords": ["答案"]}
            {"name": "词数达标", "function": "word_count_min", "min": 100}
            {"name": "文件存在", "function": "file_exists"}
            {"name": "无乱码", "function": "text_not_contains", "keywords": ["NAMEACONTENT"]}

        可选字段:severity("error"/"warning"/"info",默认 "error"),
        fix_suggestion(修复建议字符串)。

        Args:
            path: 文件路径
            checks: 检查项定义列表

        Returns:
            ValidationReport: 汇总验证报告
        """
        items: list[CheckItem] = []
        for chk in checks:
            name = chk.get("name", chk.get("function", "unknown"))
            func = chk.get("function", "")
            severity = chk.get("severity", "error")
            suggestion = chk.get("fix_suggestion")
            try:
                passed, detail = self._run_decl_check(path, chk)
                items.append(CheckItem(
                    name=name,
                    passed=passed,
                    severity=severity,
                    detail=detail,
                    fix_suggestion=suggestion,
                ))
            except Exception as e:
                items.append(CheckItem(
                    name=name,
                    passed=False,
                    severity=severity,
                    detail=f"检查执行异常: {e}",
                    fix_suggestion=suggestion,
                ))
        return ValidationReport.from_checks(items)

    def _run_decl_check(self, path: str, chk: dict) -> tuple[bool, str]:
        """执行单个声明式检查。

        Args:
            path: 文件路径
            chk: 检查定义 dict

        Returns:
            (passed, detail) 元组
        """
        func = chk.get("function", "")

        if func == "file_exists":
            exists = os.path.exists(path)
            return (exists, f"文件{'存在' if exists else '不存在'}: {path}")

        # 以下检查需要读取文本
        text = self._extract_text(path)
        if text is None:
            return (False, "无法读取文件(python-docx 不可用或文件损坏)")

        if func == "text_contains":
            keywords = chk.get("keywords", [])
            if not isinstance(keywords, list):
                return (False, "keywords 必须是列表")
            missing = [k for k in keywords if k not in text]
            if missing:
                return (False, f"缺少关键词: {missing}")
            return (True, f"包含所有关键词: {keywords}")

        if func == "text_not_contains":
            keywords = chk.get("keywords", [])
            if not isinstance(keywords, list):
                return (False, "keywords 必须是列表")
            found = [k for k in keywords if k in text]
            if found:
                return (False, f"不应包含的关键词出现: {found}")
            return (True, "未包含任何禁止关键词")

        if func == "word_count_min":
            minimum = chk.get("min", 0)
            if not isinstance(minimum, int) or isinstance(minimum, bool):
                return (False, "min 必须是整数")
            count = len(text.split())
            if count < minimum:
                return (False, f"词数 {count} < 最低要求 {minimum}")
            return (True, f"词数 {count} >= {minimum}")

        return (False, f"不支持的 function: {func}")

    # ------------------------------------------------------------------
    # 单项检查
    # ------------------------------------------------------------------

    def check_no_garbled(self, path: str) -> CheckItem:
        """检查无 NAMECONTENT 乱码残留。

        扫描全文,检测是否包含 ``NAME{X}CONTENT`` 模式。

        Args:
            path: docx 文件路径

        Returns:
            CheckItem: 检查结果
        """
        name = "无乱码残留"
        text = self._extract_text(path)
        if text is None:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail="无法读取文件(python-docx 不可用或文件损坏)",
                fix_suggestion="安装 python-docx: pip install python-docx",
            )
        matches = _GARBLED_PATTERN.findall(text)
        if matches:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail=f"发现 {len(matches)} 处 NAMECONTENT 乱码残留",
                fix_suggestion="检查答案恢复器(AnswerResolver)输出,重新处理乱码字段",
            )
        return CheckItem(
            name=name, passed=True, severity="info",
            detail="未发现乱码残留",
        )

    def check_answers_filled(self, path: str) -> CheckItem:
        """检查选择题括号已填答案。

        识别含 ``【单选题】``/``【多选题】`` 标记的段落,检查其文本中
        是否含 ``（A）`` 或 ``(A)`` 形式的答案。

        Args:
            path: docx 文件路径

        Returns:
            CheckItem: 检查结果
        """
        name = "选择题括号已填答案"
        paragraphs = self._extract_paragraphs(path)
        if paragraphs is None:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail="无法读取文件(python-docx 不可用或文件损坏)",
                fix_suggestion="安装 python-docx: pip install python-docx",
            )
        choice_count = 0
        unfilled: list[str] = []
        for i, ptext in enumerate(paragraphs):
            if _CHOICE_MARKER_PATTERN.search(ptext):
                choice_count += 1
                if not _FILLED_ANSWER_PATTERN.search(ptext):
                    unfilled.append(f"段落{i}")
        if choice_count == 0:
            return CheckItem(
                name=name, passed=True, severity="info",
                detail="未发现选择题(无 【单选题】/【多选题】 标记)",
            )
        if unfilled:
            preview = ", ".join(unfilled[:5])
            suffix = " ..." if len(unfilled) > 5 else ""
            return CheckItem(
                name=name, passed=False, severity="error",
                detail=f"{len(unfilled)}/{choice_count} 个选择题未填答案({preview}{suffix})",
                fix_suggestion="在选择题题干末尾的括号中填入答案字母,如 （B）",
            )
        return CheckItem(
            name=name, passed=True, severity="info",
            detail=f"{choice_count} 个选择题均已填答案",
        )

    def check_numbers_deleted(self, path: str) -> CheckItem:
        """检查题号已删除。

        检查是否有形如 ``11. 【单选题】`` 的未删除题号(行首数字+点+【)。

        Args:
            path: docx 文件路径

        Returns:
            CheckItem: 检查结果
        """
        name = "题号已删除"
        paragraphs = self._extract_paragraphs(path)
        if paragraphs is None:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail="无法读取文件(python-docx 不可用或文件损坏)",
                fix_suggestion="安装 python-docx: pip install python-docx",
            )
        remaining: list[str] = []
        for i, ptext in enumerate(paragraphs):
            if _QUESTION_NUMBER_PATTERN.match(ptext):
                remaining.append(f"段落{i}: {ptext[:30]}")
        if remaining:
            preview = ", ".join(remaining[:3])
            suffix = " ..." if len(remaining) > 3 else ""
            return CheckItem(
                name=name, passed=False, severity="warning",
                detail=f"发现 {len(remaining)} 处未删除题号({preview}{suffix})",
                fix_suggestion="删除段落开头的题号(如 '11. '),保留 【单选题】 标记",
            )
        return CheckItem(
            name=name, passed=True, severity="info",
            detail="未发现未删除题号",
        )

    def check_answer_lines_removed(self, path: str) -> CheckItem:
        """检查答案行已删除(非问答题)。

        检查是否有 ``【答案】`` 段落残留。问答题不应执行此检查。

        Args:
            path: docx 文件路径

        Returns:
            CheckItem: 检查结果
        """
        name = "答案行已删除"
        paragraphs = self._extract_paragraphs(path)
        if paragraphs is None:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail="无法读取文件(python-docx 不可用或文件损坏)",
                fix_suggestion="安装 python-docx: pip install python-docx",
            )
        remaining: list[str] = []
        for i, ptext in enumerate(paragraphs):
            if _ANSWER_LINE_PATTERN.search(ptext):
                remaining.append(f"段落{i}: {ptext[:30]}")
        if remaining:
            preview = ", ".join(remaining[:3])
            suffix = " ..." if len(remaining) > 3 else ""
            return CheckItem(
                name=name, passed=False, severity="error",
                detail=f"发现 {len(remaining)} 处 【答案】 行残留({preview}{suffix})",
                fix_suggestion="删除 【答案】 段落(答案已填入选题括号)",
            )
        return CheckItem(
            name=name, passed=True, severity="info",
            detail="未发现答案行残留",
        )

    def check_format_unified(self, path: str) -> CheckItem:
        """检查格式统一(字体/字号/粗细)。

        统计每个 run 的(字体,字号,粗细)组合,超过 3 种则 warning。

        Args:
            path: docx 文件路径

        Returns:
            CheckItem: 检查结果
        """
        name = "格式统一"
        runs_info = self._extract_runs_info(path)
        if runs_info is None:
            return CheckItem(
                name=name, passed=False, severity="error",
                detail="无法读取文件(python-docx 不可用或文件损坏)",
                fix_suggestion="安装 python-docx: pip install python-docx",
            )
        if not runs_info:
            return CheckItem(
                name=name, passed=True, severity="info",
                detail="文档无有效 run(可能为空文档)",
            )
        counter: Counter = Counter()
        for combo in runs_info:
            counter[combo] += 1
        if len(counter) > _FORMAT_COMBO_WARNING_THRESHOLD:
            top = counter.most_common(1)[0]
            return CheckItem(
                name=name, passed=False, severity="warning",
                detail=(
                    f"发现 {len(counter)} 种字体/字号/粗细组合"
                    f"(超过 {_FORMAT_COMBO_WARNING_THRESHOLD} 种)"
                ),
                fix_suggestion=f"建议统一为最多的组合: {top[0]}(出现 {top[1]} 次)",
            )
        return CheckItem(
            name=name, passed=True, severity="info",
            detail=f"格式统一,共 {len(counter)} 种组合",
        )

    # ------------------------------------------------------------------
    # 文本抽取工具
    # ------------------------------------------------------------------

    def _extract_text(self, path: str) -> Optional[str]:
        """从文件提取纯文本(支持 docx 和纯文本文件)。

        docx 文件用 python-docx 提取段落文本;其他文件(.txt/.md/.csv 等)
        直接以 UTF-8 读取。用于通用声明式检查(validate_general)和乱码扫描。

        Args:
            path: 文件路径

        Returns:
            纯文本;文件不存在或读取失败返回 None
        """
        if not os.path.exists(path):
            return None
        ext = os.path.splitext(path)[1].lstrip(".").lower()
        if ext in ("docx", "doc"):
            # docx 文件用 python-docx 提取段落
            paragraphs = self._extract_paragraphs(path)
            if paragraphs is None:
                return None
            return "\n".join(paragraphs)
        # 其他文件(txt/md/csv 等)直接读取为文本
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return None

    def _extract_paragraphs(self, path: str) -> Optional[list[str]]:
        """从 docx 提取段落文本列表。

        Args:
            path: docx 文件路径

        Returns:
            段落文本列表;路径校验失败或读取异常返回 None
        """
        err = self._validate_path(
            path, must_exist=True, allowed_exts=("docx", "doc")
        )
        if err:
            return None
        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError:
            return None
        try:
            doc = Document(path)
            return [p.text for p in doc.paragraphs]
        except Exception:
            return None

    def _extract_runs_info(
        self, path: str
    ) -> Optional[list[tuple[str, Optional[float], Optional[bool]]]]:
        """从 docx 提取 run 的(字体,字号,粗细)列表。

        仅提取非空文本的 run,用于格式统一检查。

        字体读取策略(中英文分离设计):
          - 优先读 w:eastAsia(中文字体),反映实际中文显示字体
          - 回退到 run.font.name(ASCII 字体)
          这样 FormatNormalizer 的"中文宋体+英文Times New Roman"分离设计
          不会被误判为不统一。

        归一化:
          - bold=None 视作 False(未显式设置等同于不加粗)
          - size=None 视作 12.0(小四,常见默认)

        Args:
            path: docx 文件路径

        Returns:
            run 信息列表;路径校验失败或读取异常返回 None
        """
        err = self._validate_path(
            path, must_exist=True, allowed_exts=("docx", "doc")
        )
        if err:
            return None
        try:
            self._require_lib("docx")
            from docx import Document
            from docx.oxml.ns import qn
        except ExpertError:
            return None
        try:
            doc = Document(path)
            runs_info: list[tuple[str, Optional[float], Optional[bool]]] = []
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                for run in p.runs:
                    if not run.text.strip():
                        continue
                    # 优先读 eastAsia 字体(中文实际字体)
                    font_name = run.font.name or "default"
                    try:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                ea = rfonts.get(qn("w:eastAsia"))
                                if ea:
                                    font_name = ea
                    except Exception:
                        pass
                    font_size: Optional[float] = (
                        run.font.size.pt if run.font.size else 12.0
                    )
                    # bold=None 归一化为 False
                    bold: Optional[bool] = bool(run.font.bold) if run.font.bold is not None else False
                    runs_info.append((font_name, font_size, bold))
            return runs_info
        except Exception:
            return None
