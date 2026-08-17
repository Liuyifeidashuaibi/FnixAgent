"""Word Expert(P2-9)。

Word 文档创建/编辑/样式/目录/合并/比较/脱敏/表格/批注/报告排版。

专家职责:
  - 创建/编辑 .docx 文档(段落/样式/表格/批注)
  - 合并/比较多文档,生成 TOC,正则脱敏
  - 精细排版报告(中文字体+英文字体分离/字号/行距/对齐/分页/分节/页码/封面)

底层依赖:
  - python-docx(可选,不可用时降级:返回 ExpertError 提示安装)

降级策略:
  - 依赖缺失 → 抛 ExpertError,上层方法捕获后返回 _failure
  - 文件 IO/PermissionError → 捕获返回 _failure,不崩溃
  - 路径穿越/扩展名不匹配 → _validate_path 拦截

排版参考(中文公文/论文/作业常用):
  - 字号对照: 初号=42pt, 小初=36pt, 一号=26pt, 小一=24pt,
    二号=22pt, 小二=18pt, 三号=16pt, 小三=15pt,
    四号=14pt, 小四=12pt, 五号=10.5pt, 小五=9pt
  - 中文字体: 宋体/黑体/楷体/仿宋, 英文字体: Times New Roman/Arial
  - 行距: 单倍=1.0, 1.5倍=1.5, 双倍=2.0
"""

# -*- coding: utf-8 -*-
# Copyright (C) 2026 FnixAgent. All rights reserved.
# Software Name: FnixAgent 智能工作台系统 V1.0
# This software and its source code are proprietary and confidential.
# Unauthorized copying, modification, distribution, or use is strictly prohibited.

from __future__ import annotations

import contextlib
import os
import re
import tempfile
from typing import Any

from fnixagent.office.base import BaseExpert, ExpertError, ExpertResult

# 中文字号 → 磅值映射
FONT_SIZE_CN: dict[str, float] = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}

# 对齐方式映射
ALIGNMENT_MAP: dict[str, int] = {
    "left": 0,
    "center": 1,
    "right": 2,
    "justify": 3,
}

class WordExpert(BaseExpert):
    """Word 文档专家。

    全部方法返回 ExpertResult,success=True 时 output 为文件路径或数据结构。

    能力边界:
      - 仅处理 .docx(python-docx 不支持 .doc 旧格式)
      - 不支持实时协作/修订追踪的细粒度操作(仅切换 trackChanges 标记)
      - 比较功能为逐段文本 diff,不含格式差异
    """

    @property
    def name(self) -> str:
        return "word"

    # ------------------------------------------------------------------
    # 创建与编辑
    # ------------------------------------------------------------------

    def create(
        self,
        output_path: str,
        title: str = "",
        content: str = "",
        paragraphs: list[dict] | None = None,
    ) -> ExpertResult:
        """创建 Word 文档。

        Args:
            output_path: 输出 .docx 路径
            title: 标题(可选)
            content: 单段正文内容(可选)
            paragraphs: 多段内容(每段 {text, style, bold, italic})

        Returns:
            ExpertResult(output=output_path)

        Raises:
            无(失败统一返回 ExpertResult(success=False))
        """
        # 输出路径校验:扩展名必须为 docx,父目录可写
        err = self._validate_path(output_path, allowed_exts=("docx",), must_exist=False)
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document()
            if title:
                doc.add_heading(title, level=0)
            if content:
                doc.add_paragraph(content)
            if paragraphs:
                for p in paragraphs:
                    text = p.get("text", "")
                    style = p.get("style", "Normal")
                    para = doc.add_paragraph(text, style=style)
                    if p.get("bold") or p.get("italic"):
                        for run in para.runs:
                            run.bold = p.get("bold", False)
                            run.italic = p.get("italic", False)
            doc.save(output_path)
            return self._success(output_path, paragraph_count=len(paragraphs or []))
        except (OSError, PermissionError) as e:
            return self._failure(f"create Word IO failed: {e}")
        except Exception as e:
            return self._failure(f"create Word failed: {e}")

    def edit(
        self,
        file_path: str,
        edits: list[dict],
        output_path: str | None = None,
    ) -> ExpertResult:
        """编辑现有 Word 文档(查找替换/插入段落/删除段落)。

        Args:
            file_path: 输入 .docx 路径
            edits: 编辑操作列表,每项 {op: "replace"|"insert"|"delete", ...}
                   - replace: {op, find, replace}
                   - insert:  {op, after_text, paragraph: {text, style}}
                   - delete:  {op, find}
            output_path: 输出路径(None 覆盖原文件)

        Returns:
            ExpertResult(output=save_path)
        """
        # 输入路径校验:必须存在且为 docx
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)
        if output_path:
            err = self._validate_path(output_path, allowed_exts=("docx",))
            if err:
                return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            for edit in edits:
                op = edit.get("op")
                if op == "replace":
                    self._replace_in_doc(doc, edit.get("find", ""), edit.get("replace", ""))
                elif op == "insert":
                    self._insert_after(doc, edit.get("after_text", ""), edit.get("paragraph", {}))
                elif op == "delete":
                    self._delete_paragraph(doc, edit.get("find", ""))
            save_path = output_path or file_path
            doc.save(save_path)
            return self._success(save_path, edits_applied=len(edits))
        except (OSError, PermissionError) as e:
            return self._failure(f"edit Word IO failed: {e}")
        except Exception as e:
            return self._failure(f"edit Word failed: {e}")

    def apply_style(
        self,
        file_path: str,
        style_name: str,
        target: str = "all",
        output_path: str | None = None,
    ) -> ExpertResult:
        """应用样式(标题/正文/强调 等)。

        Args:
            file_path: 输入 .docx 路径
            style_name: 样式名(需已存在于文档 styles 中)
            target: "all" 或匹配文本(段落包含该文本则应用)
            output_path: 输出路径;None 覆盖原文件

        Returns:
            ExpertResult(output=save_path)
        """
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)
        err = self._validate_string(style_name, "style_name")
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if target == "all" or target in para.text:
                    para.style = doc.styles[style_name]
            save_path = output_path or file_path
            doc.save(save_path)
            return self._success(save_path, style=style_name)
        except (OSError, PermissionError) as e:
            return self._failure(f"apply_style IO failed: {e}")
        except KeyError as e:
            return self._failure(f"style not found: {e}")
        except Exception as e:
            return self._failure(f"apply_style failed: {e}")

    def generate_toc(self, file_path: str, output_path: str | None = None) -> ExpertResult:
        """生成目录(TOC)。

        Args:
            file_path: 输入 .docx 路径
            output_path: 输出路径;None 覆盖原文件

        Returns:
            ExpertResult(output=save_path)

        Note:
            TOC 域插入后,用户首次在 Word 中打开按 F9 才会渲染。
        """
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            # 插入 TOC 域(用户首次打开 Word 时按 F9 更新)
            para = doc.add_paragraph()
            run = para.add_run()
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), "begin")
            run._r.append(fldChar)
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
            run._r.append(instrText)
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar2)
            save_path = output_path or file_path
            doc.save(save_path)
            return self._success(save_path)
        except (OSError, PermissionError) as e:
            return self._failure(f"generate_toc IO failed: {e}")
        except Exception as e:
            return self._failure(f"generate_toc failed: {e}")

    # ------------------------------------------------------------------
    # 合并与比较
    # ------------------------------------------------------------------

    def merge(self, file_paths: list[str], output_path: str) -> ExpertResult:
        """合并多个 Word 文档为一个。

        Args:
            file_paths: 源 .docx 路径列表(至少 1 个)
            output_path: 输出 .docx 路径

        Returns:
            ExpertResult(output=output_path, merged_count=N)
        """
        if not file_paths:
            return self._failure("file_paths is empty")
        for fp in file_paths:
            err = self._validate_path(fp, must_exist=True, allowed_exts=("docx",))
            if err:
                return self._failure(err)
        err = self._validate_path(output_path, allowed_exts=("docx",))
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            merged = Document(file_paths[0])
            for path in file_paths[1:]:
                sub = Document(path)
                for para in sub.paragraphs:
                    merged.add_paragraph(para.text)
            merged.save(output_path)
            return self._success(output_path, merged_count=len(file_paths))
        except (OSError, PermissionError) as e:
            return self._failure(f"merge IO failed: {e}")
        except Exception as e:
            return self._failure(f"merge failed: {e}")

    def compare(
        self,
        file_a: str,
        file_b: str,
        output_path: str | None = None,
    ) -> ExpertResult:
        """比较两个 Word 文档(返回差异列表)。

        简化实现:逐段对比文本,标记 added/removed/modified。

        Args:
            file_a: 文档 A 路径
            file_b: 文档 B 路径
            output_path: 未使用(保留以兼容签名)

        Returns:
            ExpertResult(output=[{op, text/old/new}, ...])
        """
        for fp in (file_a, file_b):
            err = self._validate_path(fp, must_exist=True, allowed_exts=("docx",))
            if err:
                return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc_a = Document(file_a)
            doc_b = Document(file_b)
            texts_a = [p.text for p in doc_a.paragraphs]
            texts_b = [p.text for p in doc_b.paragraphs]
            diffs = self._diff_lists(texts_a, texts_b)
            return self._success(diffs, file_a=file_a, file_b=file_b)
        except (OSError, PermissionError) as e:
            return self._failure(f"compare IO failed: {e}")
        except Exception as e:
            return self._failure(f"compare failed: {e}")

    # ------------------------------------------------------------------
    # 脱敏/表格/批注
    # ------------------------------------------------------------------

    def redact(
        self,
        file_path: str,
        patterns: list[str],
        replacement: str = "***",
        output_path: str | None = None,
    ) -> ExpertResult:
        """脱敏(按正则模式替换敏感信息)。

        Args:
            file_path: 输入 .docx 路径
            patterns: 正则模式列表
            replacement: 替换文本
            output_path: 输出路径;None 覆盖原文件

        Returns:
            ExpertResult(output=save_path, redact_count=N)
        """
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)
        if not patterns:
            return self._failure("patterns is empty")

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            redact_count = 0
            for para in doc.paragraphs:
                for pattern in patterns:
                    new_text, n = re.subn(pattern, replacement, para.text)
                    if n > 0:
                        redact_count += n
                        # 简化:整段替换(精确替换需操作 runs,这里保留降级实现)
                        for run in para.runs:
                            run.text = re.sub(pattern, replacement, run.text)
            save_path = output_path or file_path
            doc.save(save_path)
            return self._success(save_path, redact_count=redact_count)
        except re.error as e:
            return self._failure(f"invalid regex pattern: {e}")
        except (OSError, PermissionError) as e:
            return self._failure(f"redact IO failed: {e}")
        except Exception as e:
            return self._failure(f"redact failed: {e}")

    def extract_tables(self, file_path: str) -> ExpertResult:
        """提取 Word 文档中的全部表格。

        Args:
            file_path: 输入 .docx 路径

        Returns:
            ExpertResult(output=[[row_cells], ...], table_count=N)
        """
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append(rows)
            return self._success(tables, table_count=len(tables))
        except (OSError, PermissionError) as e:
            return self._failure(f"extract_tables IO failed: {e}")
        except Exception as e:
            return self._failure(f"extract_tables failed: {e}")

    def track_changes(
        self,
        file_path: str,
        enable: bool = True,
        output_path: str | None = None,
    ) -> ExpertResult:
        """启用/关闭修订模式(track changes)。

        简化实现:在 settings.xml 中写 trackChanges 标记。

        Args:
            file_path: 输入 .docx 路径
            enable: True 启用,False 关闭
            output_path: 输出路径;None 覆盖原文件

        Returns:
            ExpertResult(output=save_path, track_changes=bool)
        """
        err = self._validate_path(file_path, must_exist=True, allowed_exts=("docx",))
        if err:
            return self._failure(err)

        try:
            self._require_lib("docx")
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document(file_path)
            settings = doc.settings.element
            existing = settings.find(qn("w:trackChanges"))
            if enable and existing is None:
                el = OxmlElement("w:trackChanges")
                settings.append(el)
            elif not enable and existing is not None:
                settings.remove(existing)
            save_path = output_path or file_path
            doc.save(save_path)
            return self._success(save_path, track_changes=enable)
        except (OSError, PermissionError) as e:
            return self._failure(f"track_changes IO failed: {e}")
        except Exception as e:
            return self._failure(f"track_changes failed: {e}")

    # ------------------------------------------------------------------
    # 精细排版报告
    # ------------------------------------------------------------------

    def create_report(
        self,
        output_path: str,
        *,
        cover: dict[str, str],
        sections: list[dict[str, Any]],
        toc: bool = True,
        page_margin: tuple[float, float, float, float] = (2.54, 2.54, 3.17, 3.17),
        max_words: int = 0,
    ) -> ExpertResult:
        """创建精细排版的学术报告文档(封面+目录+正文+页码)。

        排版规范(中文论文/作业常用):
          - 封面:标题居中,信息项标签加粗+下划线填空
          - 目录页:无页码
          - 正文页:页码从第1页开始,底部居中
          - 一级标题:黑体,四号(14pt),加粗
          - 二级标题:宋体,四号(14pt),加粗
          - 三级标题:宋体,小四(12pt),加粗
          - 正文:中文宋体/英文 Times New Roman,小四(12pt),1.5倍行距
          - 图表按章节编号(如图1-1、表2-1)

        Args:
            output_path: 输出 .docx 路径
            cover: 封面信息字典,支持字段:
                title: 主标题(默认"生涯发展报告")
                subtitle: 副标题
                college: 学院
                name: 姓名
                major_class: 专业班级
                student_id: 学号
                teacher: 任课教师
                score: 成绩(默认空)
                semester: 学期(默认"2025-2026学年第2学期")
            sections: 正文章节列表,每章为 dict:
                {
                  "title": "一级标题文本",
                  "content": [ ... ],  # 段落/子标题/图表列表
                }
                content 每项为 dict:
                  - {"type": "h2", "text": "二级标题"}
                  - {"type": "h3", "text": "三级标题"}
                  - {"type": "p", "text": "正文段落"}
                  - {"type": "image", "path": "图片路径", "caption": "图注", "width": 5.5}
                  - {"type": "table", "headers": [...], "rows": [[...]], "caption": "表注"}
            toc: 是否插入目录页
            page_margin: 页边距(英寸) 顺序:上下左右
            max_words: 正文字数上限(0=不限);超限时在 metadata 中标记 warning

        Returns:
            ExpertResult(output=output_path, metadata={
                word_count: 正文字数,
                figure_count: 图表数,
                table_count: 表格数,
                warning: 超限警告(如有)
            })
        """
        err = self._validate_path(output_path, allowed_exts=("docx",), must_exist=False)
        if err:
            return self._failure(err)
        if not isinstance(sections, list):
            return self._failure("sections must be a list")

        try:
            self._require_lib("docx")
            from docx import Document
            from docx.enum.section import WD_SECTION
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls, qn
            from docx.shared import Cm, Inches, Pt
        except ExpertError as e:
            return self._failure(str(e))

        try:
            doc = Document()

            # ---- 页面设置 ----
            sec = doc.sections[0]
            sec.top_margin = Cm(page_margin[0])
            sec.bottom_margin = Cm(page_margin[1])
            sec.left_margin = Cm(page_margin[2])
            sec.right_margin = Cm(page_margin[3])

            # ---- 默认正文样式(Normal):中文宋体+英文Times New Roman ----
            style_normal = doc.styles["Normal"]
            style_normal.font.name = "Times New Roman"
            style_normal.font.size = Pt(FONT_SIZE_CN["小四"])
            rpr = style_normal.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = parse_xml(
                    f"<w:rFonts {nsdecls('w')} "
                    f'w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
                    f'w:eastAsia="宋体" w:cs="Times New Roman"/>'
                )
                rpr.insert(0, rfonts)
            else:
                rfonts.set(qn("w:ascii"), "Times New Roman")
                rfonts.set(qn("w:hAnsi"), "Times New Roman")
                rfonts.set(qn("w:eastAsia"), "宋体")
                rfonts.set(qn("w:cs"), "Times New Roman")
            pf = style_normal.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Pt(FONT_SIZE_CN["小四"] * 2)

            # ---- 封面节显式关闭页码 ----
            self._add_page_number(sec, enable=False)

            # ---- 配置标题样式 ----
            self._configure_heading_style(
                doc, "Heading 1", "黑体", FONT_SIZE_CN["四号"], True, True
            )
            self._configure_heading_style(
                doc, "Heading 2", "宋体", FONT_SIZE_CN["四号"], True, True
            )
            self._configure_heading_style(
                doc, "Heading 3", "宋体", FONT_SIZE_CN["小四"], True, True
            )

            # ---- 封面 ----
            self._build_cover(doc, cover)

            # ---- 目录节(无页码) ----
            if toc:
                sec_toc = doc.add_section(WD_SECTION.NEW_PAGE)
                sec_toc.top_margin = Cm(page_margin[0])
                sec_toc.bottom_margin = Cm(page_margin[1])
                sec_toc.left_margin = Cm(page_margin[2])
                sec_toc.right_margin = Cm(page_margin[3])
                self._add_page_number(sec_toc, enable=False)
                toc_title = doc.add_paragraph()
                toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = toc_title.add_run("目  录")
                self._set_run_font(run, "黑体", FONT_SIZE_CN["三号"], bold=True)
                self._insert_toc_field(doc)

            # ---- 正文节(页码从1开始) ----
            sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
            sec_body.top_margin = Cm(page_margin[0])
            sec_body.bottom_margin = Cm(page_margin[1])
            sec_body.left_margin = Cm(page_margin[2])
            sec_body.right_margin = Cm(page_margin[3])
            self._add_page_number(sec_body, enable=True, start=1)

            # ---- 写入正文章节 ----
            word_count = 0
            figure_count = 0
            table_count = 0
            chapter_idx = 0

            for sec_data in sections:
                chapter_idx += 1
                title_text = sec_data.get("title", "")
                if title_text:
                    h1 = doc.add_heading(level=1)
                    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = h1.add_run(title_text)
                    self._set_run_font(run, "黑体", FONT_SIZE_CN["四号"], bold=True)
                    h1.paragraph_format.first_line_indent = Pt(0)
                    word_count += len(title_text)

                fig_idx = 0
                tbl_idx = 0

                for item in sec_data.get("content", []):
                    itype = item.get("type", "p")
                    if itype == "h2":
                        h2 = doc.add_heading(level=2)
                        run = h2.add_run(item.get("text", ""))
                        self._set_run_font(run, "宋体", FONT_SIZE_CN["四号"], bold=True)
                        h2.paragraph_format.first_line_indent = Pt(0)
                        word_count += len(item.get("text", ""))
                    elif itype == "h3":
                        h3 = doc.add_heading(level=3)
                        run = h3.add_run(item.get("text", ""))
                        self._set_run_font(run, "宋体", FONT_SIZE_CN["小四"], bold=True)
                        h3.paragraph_format.first_line_indent = Pt(0)
                        word_count += len(item.get("text", ""))
                    elif itype == "p":
                        p = doc.add_paragraph()
                        text = item.get("text", "")
                        self._add_mixed_run(p, text)
                        word_count += len(text)
                    elif itype == "image":
                        fig_idx += 1
                        figure_count += 1
                        img_path = item.get("path", "")
                        caption = item.get("caption", "")
                        width_inch = item.get("width", 5.5)
                        recolor = item.get("recolor")  # 着色目标(如 "#FF0000")
                        if img_path and os.path.exists(img_path):
                            # 图片预处理:CMYK→sRGB + Recolor 着色
                            safe_img = self._ensure_srgb_image(img_path, recolor=recolor)
                            pic_para = doc.add_paragraph()
                            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pic_para.paragraph_format.first_line_indent = Pt(0)
                            run_img = pic_para.add_run()
                            run_img.add_picture(safe_img, width=Inches(width_inch))
                            # 清理临时文件(若生成了转换后的副本)
                            if safe_img != img_path:
                                with contextlib.suppress(OSError):
                                    os.remove(safe_img)
                        cap_para = doc.add_paragraph()
                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_para.paragraph_format.first_line_indent = Pt(0)
                        cap_text = (
                            f"图{chapter_idx}-{fig_idx} {caption}"
                            if caption
                            else f"图{chapter_idx}-{fig_idx}"
                        )
                        run_cap = cap_para.add_run(cap_text)
                        self._set_run_font(run_cap, "宋体", FONT_SIZE_CN["五号"], bold=False)
                        word_count += len(caption)
                    elif itype == "table":
                        tbl_idx += 1
                        table_count += 1
                        headers = item.get("headers", [])
                        rows = item.get("rows", [])
                        caption = item.get("caption", "")
                        cap_para = doc.add_paragraph()
                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_para.paragraph_format.first_line_indent = Pt(0)
                        cap_text = (
                            f"表{chapter_idx}-{tbl_idx} {caption}"
                            if caption
                            else f"表{chapter_idx}-{tbl_idx}"
                        )
                        run_cap = cap_para.add_run(cap_text)
                        self._set_run_font(run_cap, "宋体", FONT_SIZE_CN["五号"], bold=True)
                        word_count += len(caption)
                        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                        tbl.style = "Table Grid"
                        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for ci, h in enumerate(headers):
                            cell = tbl.rows[0].cells[ci]
                            cell.text = ""
                            cp = cell.paragraphs[0]
                            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_h = cp.add_run(str(h))
                            self._set_run_font(run_h, "宋体", FONT_SIZE_CN["五号"], bold=True)
                        for ri, row_data in enumerate(rows):
                            for ci, val in enumerate(row_data):
                                if ci >= len(headers):
                                    break
                                cell = tbl.rows[ri + 1].cells[ci]
                                cell.text = ""
                                cp = cell.paragraphs[0]
                                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run_c = cp.add_run(str(val))
                                self._set_run_font(run_c, "宋体", FONT_SIZE_CN["五号"], bold=False)
                        doc.add_paragraph()

            doc.save(output_path)

            warning = ""
            if max_words > 0 and word_count > max_words:
                warning = f"word_count {word_count} exceeds max_words {max_words}"

            return self._success(
                output_path,
                word_count=word_count,
                figure_count=figure_count,
                table_count=table_count,
                warning=warning,
            )
        except (OSError, PermissionError) as e:
            return self._failure(f"create_report IO failed: {e}")
        except Exception as e:
            return self._failure(f"create_report failed: {e}")

    def _configure_heading_style(
        self, doc, style_name: str, cn_font: str, size_pt: float, bold: bool, indent_zero: bool
    ) -> None:
        """配置标题样式:中文字体/英文字体/字号/加粗/行距。"""
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.shared import Pt

        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(
                f"<w:rFonts {nsdecls('w')} "
                f'w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
                f'w:eastAsia="{cn_font}" w:cs="Times New Roman"/>'
            )
            rpr.insert(0, rfonts)
        else:
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), cn_font)
            rfonts.set(qn("w:cs"), "Times New Roman")
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        if indent_zero:
            pf.first_line_indent = Pt(0)

    @staticmethod
    def _set_run_font(run, cn_font: str, size_pt: float, bold: bool = False) -> None:
        """设置 run 的字体(中文+英文分离)、字号、加粗。"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.shared import Pt

        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(
                f"<w:rFonts {nsdecls('w')} "
                f'w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
                f'w:eastAsia="{cn_font}" w:cs="Times New Roman"/>'
            )
            rpr.insert(0, rfonts)
        else:
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), cn_font)
            rfonts.set(qn("w:cs"), "Times New Roman")

    def _add_mixed_run(self, paragraph, text: str) -> None:
        """向段落添加文本,自动分离中文(宋体)与英文/数字(Times New Roman)。

        规则:连续的中文字符/中文标点用宋体(ascii/hAnsi也回退宋体);
        连续的 ASCII 字母/数字/半角标点用 Times New Roman;
        未匹配的剩余字符默认宋体。
        """
        import re

        from docx.oxml.ns import qn

        pattern = re.compile(
            r"([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+)"
            r"|([a-zA-Z0-9\s\.\,\;\:\!\?\(\)\[\]\{\}\<\>\/\\\|\-\+\=\*\&\^\%\$\#\@\!\`\~]+)"
        )
        pos = 0
        for m in pattern.finditer(text):
            start, end = m.span()
            if start > pos:
                mid = text[pos:start]
                run = paragraph.add_run(mid)
                self._set_run_font(run, "宋体", FONT_SIZE_CN["小四"])
            cn_part = m.group(1)
            en_part = m.group(2)
            if cn_part:
                run = paragraph.add_run(cn_part)
                self._set_run_font(run, "宋体", FONT_SIZE_CN["小四"])
            elif en_part:
                run = paragraph.add_run(en_part)
                self._set_run_font(run, "宋体", FONT_SIZE_CN["小四"])
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    rfonts.set(qn("w:ascii"), "Times New Roman")
                    rfonts.set(qn("w:hAnsi"), "Times New Roman")
                    rfonts.set(qn("w:cs"), "Times New Roman")
                    rfonts.set(qn("w:eastAsia"), "宋体")
            pos = end
        if pos < len(text):
            tail = text[pos:]
            run = paragraph.add_run(tail)
            self._set_run_font(run, "宋体", FONT_SIZE_CN["小四"])

    def _build_cover(self, doc, cover: dict[str, str]) -> None:
        """构建封面页。"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title = cover.get("title", "生涯发展报告")
        subtitle = cover.get("subtitle", "《大学生就业指导》结课作业")
        semester = cover.get("semester", "2025-2026学年第2学期")

        cover_info_size = FONT_SIZE_CN["小三"]

        for _ in range(2):
            doc.add_paragraph()

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(title)
        self._set_run_font(run_t, "黑体", FONT_SIZE_CN["一号"], bold=True)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_s = p_sub.add_run(subtitle)
        self._set_run_font(run_s, "宋体", FONT_SIZE_CN["二号"], bold=True)

        for _ in range(6):
            doc.add_paragraph()

        info_items = [
            ("学    院：", cover.get("college", "")),
            ("姓    名：", cover.get("name", "")),
            ("专业班级：", cover.get("major_class", "")),
            ("学    号：", cover.get("student_id", "")),
            ("任课教师：", cover.get("teacher", "")),
        ]
        for label, value in info_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_l = p.add_run(label)
            self._set_run_font(run_l, "宋体", cover_info_size, bold=True)
            run_v = p.add_run(value + "                    ")
            self._set_run_font(run_v, "宋体", cover_info_size, bold=False)
            run_v.font.underline = True

        for _ in range(2):
            doc.add_paragraph()

        p_score = doc.add_paragraph()
        p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sl = p_score.add_run("成    绩：")
        self._set_run_font(run_sl, "宋体", cover_info_size, bold=True)
        run_sv = p_score.add_run(cover.get("score", "") + "                    ")
        self._set_run_font(run_sv, "宋体", cover_info_size, bold=False)
        run_sv.font.underline = True

        for _ in range(3):
            doc.add_paragraph()

        p_sem = doc.add_paragraph()
        p_sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sem = p_sem.add_run(semester)
        self._set_run_font(run_sem, "宋体", cover_info_size, bold=True)

    def _add_page_number(self, section, enable: bool = True, start: int = 1) -> None:
        """为节添加/移除页码(页脚居中)。

        Args:
            section: docx Section 对象
            enable: True 添加页码,False 清除页码
            start: 起始页码
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        footer = section.footer
        footer.is_linked_to_previous = False

        if enable:
            sectPr = section._sectPr
            pgNumType = sectPr.find(qn("w:pgNumType"))
            if pgNumType is None:
                pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="{start}"/>')
                sectPr.append(pgNumType)
            else:
                pgNumType.set(qn("w:start"), str(start))

            for p in list(footer.paragraphs):
                for run in list(p.runs):
                    run._element.getparent().remove(run._element)
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            p.alignment = 1
            run = p.add_run()
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fldChar1)
            instrText = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            )
            run._r.append(instrText)
            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run._r.append(fldChar2)
            self._set_run_font(run, "宋体", FONT_SIZE_CN["小五"])
        else:
            sectPr = section._sectPr
            pgNumType = sectPr.find(qn("w:pgNumType"))
            if pgNumType is not None:
                sectPr.remove(pgNumType)
            for p in footer.paragraphs:
                for run in list(p.runs):
                    run._element.getparent().remove(run._element)

    def _insert_toc_field(self, doc) -> None:
        """在当前位置插入 TOC 域(Word 打开时按 F9 可更新)。"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        para = doc.add_paragraph()
        run = para.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">TOC \\o "1-3" \\h \\z \\u</w:instrText>'
        )
        run._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar2)
        placeholder = parse_xml(
            f'<w:t {nsdecls("w")} xml:space="preserve">请在 Word 中按 F9 更新目录</w:t>'
        )
        run._r.append(placeholder)
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar3)

    @staticmethod
    def _ensure_srgb_image(image_path: str, recolor: str | None = None) -> str:
        """图片预处理:色彩空间转换 + Recolor 着色。

        处理两个问题:
          1. CMYK→sRGB 转换:避免 LibreOffice 转 PDF 时色偏
          2. Recolor 着色:Word 对单色透明 PNG 自动应用主题色重新着色
             (如黄色校徽 Logo 被映射为红色主题色),但 LibreOffice 不应用
             此效果导致 PDF 显示源色(黄色)而非 Word 中的着色(红色)
             解决方案:主动将图片着色为目标颜色,使 docx 和 PDF 显示一致

        Args:
            image_path: 原始图片路径
            recolor: 目标着色颜色(如 "#FF0000" 红色)
                     指定后,单色透明 PNG 的主色将被替换为此颜色
                     None 则不进行着色(仅做色彩空间转换)

        Returns:
            处理后的图片路径(若无需处理则返回原路径)

        Note:
            转换后生成临时 PNG 文件,调用方负责清理。
        """
        try:
            import io as _io

            import numpy as np
            from PIL import Image, ImageCms, ImageColor
        except ImportError:
            return image_path

        try:
            img = Image.open(image_path)
            mode = img.mode
            needs_processing = False

            # ---- CMYK→sRGB 色彩空间转换 ----
            if mode == "CMYK":
                src_icc = img.info.get("icc_profile")
                if src_icc:
                    try:
                        srgb_profile = ImageCms.createProfile("sRGB")
                        src_profile = ImageCms.ImageCmsProfile(_io.BytesIO(src_icc))
                        img = ImageCms.profileToProfile(
                            img, src_profile, srgb_profile, outputMode="RGB"
                        )
                    except Exception:
                        img = img.convert("RGB")
                else:
                    img = img.convert("RGB")
                needs_processing = True
            elif mode == "P":
                img = img.convert("RGBA")
                mode = "RGBA"
                needs_processing = True

            # ---- Recolor 着色:主动应用颜色映射 ----
            # 检测单色/少色透明 PNG,将主色替换为目标颜色
            # 同时扁平化透明背景为白色,防止 Word 二次 Recolor
            if mode == "RGBA" or (recolor and mode in ("RGB", "RGBA")):
                alpha = img.split()[3] if mode == "RGBA" else None
                has_transparency = False
                if alpha is not None:
                    alpha_arr = np.array(alpha)
                    has_transparency = (alpha_arr < 250).any()

                rgb_img = img.convert("RGB")
                arr = np.array(rgb_img).reshape(-1, 3)

                # 确定不透明像素
                if alpha is not None:
                    opaque_mask = alpha_arr.flatten() >= 250
                    opaque_pixels = arr[opaque_mask]
                else:
                    opaque_mask = np.ones(len(arr), dtype=bool)
                    opaque_pixels = arr

                if len(opaque_pixels) > 0:
                    quantized = (opaque_pixels // 32) * 32
                    unique_colors = len(set(map(tuple, quantized)))

                    # 单色或少色图片 + 透明背景 → Recolor 候选
                    is_recolor_candidate = unique_colors <= 3 and has_transparency

                    if is_recolor_candidate and recolor:
                        # 主动着色:将所有不透明像素替换为目标颜色
                        target_rgb = ImageColor.getrgb(recolor)
                        # 创建着色后的 RGB 图
                        colored_arr = np.full_like(arr, target_rgb)
                        # 只替换不透明像素
                        colored_arr[opaque_mask] = target_rgb
                        colored_arr = colored_arr.reshape(rgb_img.size[1], rgb_img.size[0], 3)
                        img = Image.fromarray(colored_arr, "RGB")
                        needs_processing = True
                    elif is_recolor_candidate:
                        # 无 recolor 参数:扁平化透明背景为白色
                        # 防止 Word 自动 Recolor 导致 docx/PDF 不一致
                        background = Image.new("RGB", img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[3])
                        img = background
                        needs_processing = True

            if mode in ("L", "LA") and mode != "RGBA":
                img = img.convert("RGB")
                needs_processing = True

            if not needs_processing and mode in ("RGB",):
                return image_path

            # 保存为临时 PNG
            fd, tmp_path = tempfile.mkstemp(suffix=".png", prefix="img_")
            os.close(fd)
            if img.mode != "RGB":
                img = img.convert("RGB")
            img.save(tmp_path, "PNG")
            return tmp_path
        except Exception:
            return image_path

    # ------------------------------------------------------------------
    # 内部工具
    # ------------------------------------------------------------------

    def _replace_in_doc(self, doc, find: str, replace: str) -> None:
        """在文档全部段落中查找替换。"""
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)

    def _insert_after(self, doc, after_text: str, paragraph: dict) -> None:
        """在指定文本后插入新段落。"""
        for para in doc.paragraphs:
            if after_text in para.text:
                doc.add_paragraph(
                    paragraph.get("text", ""),
                    paragraph.get("style", "Normal"),
                )
                # 移动到目标位置后(简化:追加到末尾)
                return

    def _delete_paragraph(self, doc, find: str) -> None:
        """删除包含指定文本的段落。"""
        for para in list(doc.paragraphs):
            if find in para.text:
                p = para._element
                p.getparent().remove(p)

    def _diff_lists(self, a: list[str], b: list[str]) -> list[dict]:
        """简易列表差异(逐项对比)。"""
        diffs = []
        max_len = max(len(a), len(b))
        for i in range(max_len):
            ta = a[i] if i < len(a) else None
            tb = b[i] if i < len(b) else None
            if ta is None and tb is not None:
                diffs.append({"op": "added", "text": tb})
            elif ta is not None and tb is None:
                diffs.append({"op": "removed", "text": ta})
            elif ta != tb:
                diffs.append({"op": "modified", "old": ta, "new": tb})
        return diffs

    # ------------------------------------------------------------------
    # 表格格式化(借鉴 Office-Word-MCP-Server/core/tables.py)
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell_shading(cell, fill_color: str, pattern: str = "clear") -> bool:
        """设置单元格底纹颜色。

        Args:
            cell: python-docx 的 _Cell 对象
            fill_color: 填充颜色(hex 字符串,如 "FF0000" 或 "#FF0000")
            pattern: 底纹图案("clear"/"solid"/"pct10" 等)

        Returns:
            是否成功
        """
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls, qn

            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:shd"))
            if existing is not None:
                tc_pr.remove(existing)

            color = fill_color.lstrip("#").upper()
            if len(color) != 6:
                return False
            shd_xml = f'<w:shd {nsdecls("w")} w:val="{pattern}" w:color="auto" w:fill="{color}"/>'
            tc_pr.append(parse_xml(shd_xml))
            return True
        except Exception:
            return False

    @staticmethod
    def _set_cell_border(cell, val: str = "single", sz: str = "4", color: str = "000000") -> bool:
        """设置单元格四边边框。

        Args:
            cell: python-docx 的 _Cell 对象
            val: 边框样式("single"/"double"/"thick"/"nil")
            sz: 边框粗细(1/8 pt 为单位,如 "4"=0.5pt)
            color: 边框颜色(hex 字符串)

        Returns:
            是否成功
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml.shared import OxmlElement

            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)

            for edge in ("top", "left", "bottom", "right"):
                existing = tc_borders.find(qn(f"w:{edge}"))
                if existing is not None:
                    tc_borders.remove(existing)
                elem = OxmlElement(f"w:{edge}")
                elem.set(qn("w:val"), val)
                elem.set(qn("w:sz"), sz)
                elem.set(qn("w:space"), "0")
                elem.set(qn("w:color"), color)
                tc_borders.append(elem)
            return True
        except Exception:
            return False

    @staticmethod
    def _set_cell_alignment(cell, horizontal: str = "left", vertical: str = "top") -> bool:
        """设置单元格文本对齐方式。

        Args:
            cell: python-docx 的 _Cell 对象
            horizontal: 水平对齐("left"/"center"/"right"/"justify")
            vertical: 垂直对齐("top"/"center"/"bottom")

        Returns:
            是否成功
        """
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml.shared import OxmlElement

            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            for para in cell.paragraphs:
                para.alignment = align_map.get(horizontal, WD_ALIGN_PARAGRAPH.LEFT)

            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:vAlign"))
            if existing is not None:
                tc_pr.remove(existing)
            v_elem = OxmlElement("w:vAlign")
            v_elem.set(qn("w:val"), vertical)
            tc_pr.append(v_elem)
            return True
        except Exception:
            return False

    @staticmethod
    def _merge_cells(table, start_row: int, start_col: int, end_row: int, end_col: int) -> bool:
        """合并矩形区域内的单元格。

        Args:
            table: python-docx 的 Table 对象
            start_row: 起始行(0-based)
            start_col: 起始列(0-based)
            end_row: 结束行(0-based, 含)
            end_col: 结束列(0-based, 含)

        Returns:
            是否成功
        """
        try:
            if start_row < 0 or end_row >= len(table.rows):
                return False
            for ri in range(start_row, end_row + 1):
                if start_col < 0 or end_col >= len(table.rows[ri].cells):
                    return False
            table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
            return True
        except Exception:
            return False

    @staticmethod
    def _set_column_width(table, col_index: int, width_pt: float, width_type: str = "dxa") -> bool:
        """设置表格列宽。

        Args:
            table: python-docx 的 Table 对象
            col_index: 列索引(0-based)
            width_pt: 宽度值(磅)
            width_type: 宽度类型("dxa"=磅*20, "pct"=百分比*50, "auto")

        Returns:
            是否成功
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml.shared import OxmlElement

            if col_index < 0 or col_index >= len(table.columns):
                return False

            if width_type == "dxa":
                width_val = str(int(width_pt * 20))
            elif width_type == "pct":
                width_val = str(int(width_pt * 50))
            else:
                width_val = str(width_pt)

            for row in table.rows:
                if col_index < len(row.cells):
                    tc_pr = row.cells[col_index]._tc.get_or_add_tcPr()
                    existing = tc_pr.find(qn("w:tcW"))
                    if existing is not None:
                        tc_pr.remove(existing)
                    w_elem = OxmlElement("w:tcW")
                    w_elem.set(qn("w:w"), width_val)
                    w_elem.set(qn("w:type"), width_type)
                    tc_pr.append(w_elem)
            return True
        except Exception:
            return False

    @staticmethod
    def _apply_alternating_rows(table, color1: str = "FFFFFF", color2: str = "F2F2F2") -> bool:
        """应用交替行底纹(提升可读性)。

        Args:
            table: python-docx 的 Table 对象
            color1: 奇数行颜色(hex)
            color2: 偶数行颜色(hex)

        Returns:
            是否成功
        """
        try:
            for i, row in enumerate(table.rows):
                fill = color1 if i % 2 == 0 else color2
                for cell in row.cells:
                    WordExpert._set_cell_shading(cell, fill)
            return True
        except Exception:
            return False

    @staticmethod
    def _highlight_header(
        table,
        bg_color: str = "4472C4",
        text_color: str = "FFFFFF",
    ) -> bool:
        """高亮表头行(底纹+加粗+文字颜色)。

        Args:
            table: python-docx 的 Table 对象
            bg_color: 表头背景色(hex)
            text_color: 表头文字色(hex)

        Returns:
            是否成功
        """
        try:
            from docx.shared import RGBColor

            if not table.rows:
                return False
            for cell in table.rows[0].cells:
                WordExpert._set_cell_shading(cell, bg_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        tc = text_color.lstrip("#")
                        try:
                            run.font.color.rgb = RGBColor(
                                int(tc[0:2], 16),
                                int(tc[2:4], 16),
                                int(tc[4:6], 16),
                            )
                        except Exception:
                            pass
            return True
        except Exception:
            return False

    @staticmethod
    def _format_table_full(
        table,
        has_header: bool = True,
        border_style: str = "single",
        header_bg: str = "4472C4",
        header_text_color: str = "FFFFFF",
        alt_row_colors: tuple[str, str] | None = None,
    ) -> bool:
        """一站式表格格式化(表头+边框+交替行)。

        Args:
            table: python-docx 的 Table 对象
            has_header: 是否格式化首行为表头
            border_style: 边框样式("single"/"double"/"thick"/"nil")
            header_bg: 表头背景色
            header_text_color: 表头文字色
            alt_row_colors: 交替行颜色元组(None 则不交替)

        Returns:
            是否成功
        """
        try:
            # 边框
            for row in table.rows:
                for cell in row.cells:
                    WordExpert._set_cell_border(cell, val=border_style)

            # 表头
            if has_header:
                WordExpert._highlight_header(table, header_bg, header_text_color)

            # 交替行(跳过表头)
            if alt_row_colors:
                start = 1 if has_header else 0
                for i in range(start, len(table.rows)):
                    fill = alt_row_colors[0] if (i - start) % 2 == 0 else alt_row_colors[1]
                    for cell in table.rows[i].cells:
                        WordExpert._set_cell_shading(cell, fill)
            return True
        except Exception:
            return False
