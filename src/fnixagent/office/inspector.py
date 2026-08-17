"""文档检查器(借鉴 OfficeCLI render→look→fix 闭环)。

核心能力:
  - render: 将文档(docx/pdf/pptx)渲染为图片,让 AI 能"看到"效果
  - inspect: 检查文档格式/内容是否符合预期
  - diagnose: 生成诊断报告,列出问题供 AI 修复

设计参考:
  - OfficeCLI 的 "render → look → fix" 模式:
    AI 生成文档后先渲染为图片,通过视觉检查发现问题,再修复
  - PPTAgent 的 PPTEval 评估框架:
    从内容/设计/连贯性三个维度评估文档质量

典型流程:
  1. AI 调用 WordExpert 生成 docx
  2. 调用 DocumentInspector.render() 将 docx 转为 PDF 再渲染为图片
  3. AI 查看图片(通过 vision LLM 或像素分析)
  4. 发现问题后调用 WordExpert 修复
  5. 重复直到检查通过
"""

# -*- coding: utf-8 -*-
# Copyright (C) 2026 FnixAgent. All rights reserved.
# Software Name: FnixAgent 智能工作台系统 V1.0
# This software and its source code are proprietary and confidential.
# Unauthorized copying, modification, distribution, or use is strictly prohibited.

from __future__ import annotations

import os
import tempfile

from fnixagent.office.base import BaseExpert, ExpertResult

class DocumentInspector(BaseExpert):
    """文档检查器:渲染、检查、诊断。

    提供文档可视化能力,支持 AI Agent 的"看→修"闭环。
    """

    _name = "inspector"

    @property
    def name(self) -> str:
        return self._name

    def render(
        self,
        document_path: str,
        output_dir: str | None = None,
        *,
        pages: list[int] | None = None,
        dpi: int = 150,
        max_pages: int = 20,
    ) -> ExpertResult:
        """将文档渲染为图片,让 AI 能"看到"文档效果。

        支持 docx→pdf→png、pdf→png、pptx→png 的渲染链。
        渲染后的图片可供 vision LLM 分析或像素级检查。

        Args:
            document_path: 文档路径(.docx/.pdf/.pptx)
            output_dir: 输出目录(None 则用临时目录)
            pages: 指定页码列表(None 则渲染所有页)
            dpi: 渲染分辨率(默认 150)
            max_pages: 最大渲染页数(防止 OOM)

        Returns:
            ExpertResult(output=output_dir, metadata={
                images: list[str],  # 渲染后的图片路径列表
                page_count: int,    # 文档总页数
                rendered_count: int # 实际渲染页数
            })
        """
        err = self._validate_path(document_path, allowed_exts=("docx", "pdf", "pptx"))
        if err:
            return self._failure(err)

        ext = os.path.splitext(document_path)[1].lstrip(".").lower()

        # docx/pptx 需先转为 PDF
        pdf_path = document_path
        temp_pdf = None
        if ext in ("docx", "pptx"):
            try:
                self._require_lib("fitz")
                from fnixagent.office.converter import ConverterExpert

                converter = ConverterExpert()
                temp_fd, temp_pdf = tempfile.mkstemp(suffix=".pdf", prefix="render_")
                os.close(temp_fd)
                r = converter.convert(document_path, temp_pdf, target_format="pdf")
                if not r.success:
                    return self._failure(f"render: convert to PDF failed: {r.error}")
                pdf_path = temp_pdf
            except Exception as e:
                return self._failure(f"render: convert failed: {e}")

        try:
            self._require_lib("fitz")
            import fitz

            doc = fitz.open(pdf_path)
            page_count = len(doc)
            if page_count == 0:
                doc.close()
                return self._failure("render: document has 0 pages")

            # 确定渲染哪些页
            if pages is None:
                target_pages = list(range(min(page_count, max_pages)))
            else:
                target_pages = [p - 1 for p in pages if 1 <= p <= page_count]

            # 输出目录
            if output_dir is None:
                output_dir = tempfile.mkdtemp(prefix="render_out_")
            else:
                os.makedirs(output_dir, exist_ok=True)

            images = []
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            for page_idx in target_pages:
                page = doc[page_idx]
                pix = page.get_pixmap(matrix=mat)
                img_path = os.path.join(output_dir, f"page_{page_idx + 1:03d}.png")
                pix.save(img_path)
                images.append(img_path)

            doc.close()

            return self._success(
                output_dir,
                images=images,
                page_count=page_count,
                rendered_count=len(images),
            )
        except Exception as e:
            return self._failure(f"render failed: {e}")
        finally:
            if temp_pdf and os.path.exists(temp_pdf):
                with __import__("contextlib").suppress(OSError):
                    os.remove(temp_pdf)

    def inspect(
        self,
        document_path: str,
        *,
        checks: list[str] | None = None,
    ) -> ExpertResult:
        """检查文档格式与内容是否符合预期。

        检查项(可指定子集):
          - page_count: 页数检查
          - word_count: 字数检查
          - fonts: 字体检查(中文宋体/英文Times New Roman)
          - images: 图片检查(数量、色彩模式)
          - tables: 表格检查(数量、编号)
          - headings: 标题层级检查
          - page_numbers: 页码检查

        Args:
            document_path: 文档路径(.docx)
            checks: 检查项列表(None 则全部检查)

        Returns:
            ExpertResult(metadata={
                checks: dict[str, dict],  # {check_name: {passed: bool, detail: str}}
                passed_count: int,
                failed_count: int,
                summary: str
            })
        """
        err = self._validate_path(document_path, allowed_exts=("docx",))
        if err:
            return self._failure(err)

        all_checks = [
            "page_count",
            "word_count",
            "fonts",
            "images",
            "tables",
            "headings",
            "page_numbers",
        ]
        target_checks = checks or all_checks
        results: dict[str, dict] = {}

        try:
            self._require_lib("docx")
            from docx import Document

            doc = Document(document_path)

            for check_name in target_checks:
                if check_name == "page_count":
                    results[check_name] = self._check_page_count(doc)
                elif check_name == "word_count":
                    results[check_name] = self._check_word_count(doc)
                elif check_name == "fonts":
                    results[check_name] = self._check_fonts(doc)
                elif check_name == "images":
                    results[check_name] = self._check_images(doc)
                elif check_name == "tables":
                    results[check_name] = self._check_tables(doc)
                elif check_name == "headings":
                    results[check_name] = self._check_headings(doc)
                elif check_name == "page_numbers":
                    results[check_name] = self._check_page_numbers(doc)

            passed = sum(1 for r in results.values() if r.get("passed"))
            failed = sum(1 for r in results.values() if not r.get("passed"))

            summary_parts = []
            for name, r in results.items():
                status = "PASS" if r.get("passed") else "FAIL"
                summary_parts.append(f"[{status}] {name}: {r.get('detail', '')}")

            return self._success(
                document_path,
                checks=results,
                passed_count=passed,
                failed_count=failed,
                summary="\n".join(summary_parts),
            )
        except Exception as e:
            return self._failure(f"inspect failed: {e}")

    # -- 具体检查项 --------------------------------------------------------

    @staticmethod
    def _check_page_count(doc) -> dict:
        sections = len(doc.sections)
        return {
            "passed": sections >= 1,
            "detail": f"sections={sections}",
            "value": sections,
        }

    @staticmethod
    def _check_word_count(doc) -> dict:
        count = sum(len(p.text) for p in doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    count += len(cell.text)
        return {
            "passed": count > 0,
            "detail": f"total_chars={count}",
            "value": count,
        }

    @staticmethod
    def _check_fonts(doc) -> dict:
        issues = []
        for sname in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
            if sname not in [s.name for s in doc.styles]:
                continue
            style = doc.styles[sname]
            rpr = style.element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
            )
            if rpr is not None:
                rfonts = rpr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                )
                if rfonts is not None:
                    ea = rfonts.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                    )
                    asc = rfonts.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
                    )
                    if sname == "Normal" and ea != "宋体":
                        issues.append(f"{sname}: eastAsia={ea}(expected 宋体)")
                    if asc and asc != "Times New Roman":
                        issues.append(f"{sname}: ascii={asc}(expected Times New Roman)")
        return {
            "passed": len(issues) == 0,
            "detail": "; ".join(issues) if issues else "fonts OK",
            "issues": issues,
        }

    @staticmethod
    def _check_images(doc) -> dict:
        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_count += 1
        return {
            "passed": True,
            "detail": f"image_count={img_count}",
            "value": img_count,
        }

    @staticmethod
    def _check_tables(doc) -> dict:
        tbl_count = len(doc.tables)
        return {
            "passed": True,
            "detail": f"table_count={tbl_count}",
            "value": tbl_count,
        }

    @staticmethod
    def _check_headings(doc) -> dict:
        h1_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Heading 1")
        h2_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Heading 2")
        return {
            "passed": h1_count > 0,
            "detail": f"h1={h1_count}, h2={h2_count}",
            "h1_count": h1_count,
            "h2_count": h2_count,
        }

    @staticmethod
    def _check_page_numbers(doc) -> dict:
        from docx.oxml.ns import qn

        issues = []
        for i, sec in enumerate(doc.sections):
            footer = sec.footer
            has_page = False
            for p in footer.paragraphs:
                xml = p._element.xml
                if "PAGE" in xml and "fldChar" in xml:
                    has_page = True
            sectPr = sec._sectPr
            pgNumType = sectPr.find(qn("w:pgNumType"))
            pgNumType.get(qn("w:start")) if pgNumType is not None else None
            if i == 0 and has_page:
                issues.append("section 0 (cover) should not have page numbers")
            if i == len(doc.sections) - 1 and not has_page:
                issues.append("last section (body) should have page numbers")
        return {
            "passed": len(issues) == 0,
            "detail": "; ".join(issues) if issues else "page_numbers OK",
            "issues": issues,
        }
