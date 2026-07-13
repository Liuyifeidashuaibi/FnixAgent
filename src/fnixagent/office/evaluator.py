"""Office Agent 评测框架(借鉴 OfficeBench + SpreadsheetBench)。

设计参考:
  - OfficeBench 的声明式评测:任务 JSON 中定义 evaluation 检查项列表
  - SpreadsheetBench 的 OJ 式多测试用例 + Soft/Hard 双指标
  - 评测与任务解耦,新增任务无需改评测代码

核心概念:
  - Task: 一个办公任务(自然语言描述 + 评测项列表)
  - Check: 一个声明式检查项(function_name + args)
  - TestCase: 同一任务的不同参数变体(测试鲁棒性)
  - Result: 单次评测结果(passed/failed/detail)
  - Score: 汇总分数(soft_rate/hard_rate/total)

典型用法:
  evaluator = Evaluator()
  result = evaluator.evaluate_task(
      output_dir="path/to/output",
      checks=[
          {"function": "file_exists", "args": {"filename": "report.docx"}},
          {"function": "text_contains", "args": {"filename": "report.docx",
              "keywords": ["职业目标", "行动成效"]}},
      ]
  )
  # result.soft_rate = 1.0  # 全部通过
  # result.hard_rate = 1.0  # 硬约束通过
"""
from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass, field
from typing import Any, Callable, Optional

from fnixagent.office.base import BaseExpert, ExpertResult


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------


@dataclass
class CheckResult:
    """单个检查项的结果。

    Attributes:
        passed: 是否通过
        function: 检查函数名
        detail: 详细描述
    """

    passed: bool
    function: str
    detail: str = ""


@dataclass
class TaskResult:
    """单个任务的评测结果。

    Attributes:
        task_id: 任务 ID
        checks: 各检查项结果列表
        soft_rate: 软通过率(通过项数/总项数)
        hard_rate: 硬通过率(全部通过=1, 否则=0)
        passed_count: 通过的检查项数
        total_count: 总检查项数
    """

    task_id: str = ""
    checks: list[CheckResult] = field(default_factory=list)
    soft_rate: float = 0.0
    hard_rate: float = 0.0
    passed_count: int = 0
    total_count: int = 0

    def compute(self) -> None:
        """根据 checks 计算 soft_rate 和 hard_rate。"""
        self.total_count = len(self.checks)
        self.passed_count = sum(1 for c in self.checks if c.passed)
        if self.total_count > 0:
            self.soft_rate = self.passed_count / self.total_count
            self.hard_rate = 1.0 if self.passed_count == self.total_count else 0.0


@dataclass
class ScoreSummary:
    """评测汇总分数(借鉴 SpreadsheetBench 的 Soft/Hard 双指标)。

    Attributes:
        total_tasks: 总任务数
        passed_tasks: 硬约束通过的任务数
        soft_avg: 软通过率平均值
        hard_rate: 硬约束通过率(全通过任务占比)
        results: 各任务结果
    """

    total_tasks: int = 0
    passed_tasks: int = 0
    soft_avg: float = 0.0
    hard_rate: float = 0.0
    results: list[TaskResult] = field(default_factory=list)

    def compute(self) -> None:
        """汇总计算。"""
        self.total_tasks = len(self.results)
        self.passed_tasks = sum(1 for r in self.results if r.hard_rate == 1.0)
        if self.total_tasks > 0:
            self.soft_avg = sum(r.soft_rate for r in self.results) / self.total_tasks
            self.hard_rate = self.passed_tasks / self.total_tasks


# ---------------------------------------------------------------------------
# 评测器
# ---------------------------------------------------------------------------


class Evaluator(BaseExpert):
    """Office Agent 评测器。

    提供声明式评测函数库,支持:
      - 文件存在性检查
      - 文本内容包含/不包含检查
      - Excel 单元格值检查
      - Word 标题层级检查
      - 文件格式检查
      - 页数检查

    支持 OJ 式多测试用例评估(同一任务多个变体)。
    """

    _name = "evaluator"

    @property
    def name(self) -> str:
        return self._name

    def __init__(self) -> None:
        # 评测函数注册表(借鉴 OfficeBench 的声明式评测函数库)
        self._check_functions: dict[str, Callable] = {
            "file_exists": self._check_file_exists,
            "file_not_exists": self._check_file_not_exists,
            "text_contains": self._check_text_contains,
            "text_not_contains": self._check_text_not_contains,
            "excel_cell_value": self._check_excel_cell_value,
            "word_heading_exists": self._check_word_heading_exists,
            "word_word_count": self._check_word_count,
            "pdf_page_count": self._check_pdf_page_count,
            "file_extension": self._check_file_extension,
        }

    def evaluate_task(
        self,
        output_dir: str,
        checks: list[dict[str, Any]],
        task_id: str = "",
    ) -> ExpertResult:
        """评测单个任务。

        Args:
            output_dir: 产物目录
            checks: 检查项列表,每项含 function 和 args
            task_id: 任务 ID(用于追踪)

        Returns:
            ExpertResult(output=TaskResult, metadata=summary)
        """
        result = TaskResult(task_id=task_id)

        for check in checks:
            func_name = check.get("function", "")
            args = check.get("args", {})
            check_result = self._run_check(output_dir, func_name, args)
            result.checks.append(check_result)

        result.compute()

        return self._success(
            result,
            task_id=task_id,
            soft_rate=result.soft_rate,
            hard_rate=result.hard_rate,
            passed=result.passed_count,
            total=result.total_count,
        )

    def evaluate_batch(
        self,
        tasks: list[dict[str, Any]],
    ) -> ExpertResult:
        """批量评测多个任务(借鉴 SpreadsheetBench 的 OJ 式评估)。

        Args:
            tasks: 任务列表,每项含 output_dir, checks, task_id

        Returns:
            ExpertResult(output=ScoreSummary, metadata=summary)
        """
        summary = ScoreSummary()

        for task in tasks:
            output_dir = task.get("output_dir", "")
            checks = task.get("checks", [])
            task_id = task.get("task_id", "")
            r = self.evaluate_task(output_dir, checks, task_id)
            if r.success and r.output:
                summary.results.append(r.output)

        summary.compute()

        return self._success(
            summary,
            total_tasks=summary.total_tasks,
            passed_tasks=summary.passed_tasks,
            soft_avg=summary.soft_avg,
            hard_rate=summary.hard_rate,
        )

    # ------------------------------------------------------------------
    # 内部:检查函数执行
    # ------------------------------------------------------------------

    def _run_check(
        self, output_dir: str, func_name: str, args: dict
    ) -> CheckResult:
        """执行单个检查项。"""
        func = self._check_functions.get(func_name)
        if func is None:
            return CheckResult(
                passed=False,
                function=func_name,
                detail=f"unknown check function: {func_name}",
            )
        try:
            return func(output_dir, args)
        except Exception as e:
            return CheckResult(
                passed=False,
                function=func_name,
                detail=f"check error: {e}",
            )

    # ------------------------------------------------------------------
    # 评测函数库(借鉴 OfficeBench 的 evaluate.py)
    # ------------------------------------------------------------------

    @staticmethod
    def _check_file_exists(output_dir: str, args: dict) -> CheckResult:
        """检查文件是否存在。"""
        filename = args.get("filename", "")
        filepath = os.path.join(output_dir, filename)
        exists = os.path.exists(filepath)
        return CheckResult(
            passed=exists,
            function="file_exists",
            detail=f"{filename}: {'exists' if exists else 'NOT found'}",
        )

    @staticmethod
    def _check_file_not_exists(output_dir: str, args: dict) -> CheckResult:
        """检查文件不存在。"""
        filename = args.get("filename", "")
        filepath = os.path.join(output_dir, filename)
        not_exists = not os.path.exists(filepath)
        return CheckResult(
            passed=not_exists,
            function="file_not_exists",
            detail=f"{filename}: {'not exists' if not_exists else 'still exists'}",
        )

    @staticmethod
    def _check_text_contains(output_dir: str, args: dict) -> CheckResult:
        """检查文件是否包含所有关键词。"""
        filename = args.get("filename", "")
        keywords = args.get("keywords", [])
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=False, function="text_contains", detail=f"file not found: {filename}"
            )

        # 根据扩展名提取文本
        ext = os.path.splitext(filename)[1].lower()
        text = ""
        try:
            if ext == ".docx":
                from docx import Document
                doc = Document(filepath)
                text = "\n".join(p.text for p in doc.paragraphs)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            text += "\n" + cell.text
            elif ext == ".txt":
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()
            elif ext == ".pdf":
                import fitz
                doc = fitz.open(filepath)
                text = "\n".join(page.get_text() for page in doc)
                doc.close()
            else:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception as e:
            return CheckResult(
                passed=False, function="text_contains", detail=f"read error: {e}"
            )

        # 检查所有关键词
        missing = [kw for kw in keywords if kw not in text]
        if not missing:
            return CheckResult(
                passed=True, function="text_contains", detail=f"all {len(keywords)} keywords found"
            )
        return CheckResult(
            passed=False,
            function="text_contains",
            detail=f"missing keywords: {missing}",
        )

    @staticmethod
    def _check_text_not_contains(output_dir: str, args: dict) -> CheckResult:
        """检查文件不包含某些关键词。"""
        filename = args.get("filename", "")
        keywords = args.get("keywords", [])
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=True, function="text_not_contains", detail="file not found (vacuous truth)"
            )

        ext = os.path.splitext(filename)[1].lower()
        text = ""
        try:
            if ext == ".docx":
                from docx import Document
                doc = Document(filepath)
                text = "\n".join(p.text for p in doc.paragraphs)
            elif ext == ".txt":
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception:
            return CheckResult(
                passed=False, function="text_not_contains", detail="read error"
            )

        found = [kw for kw in keywords if kw in text]
        if not found:
            return CheckResult(
                passed=True, function="text_not_contains", detail="no forbidden keywords found"
            )
        return CheckResult(
            passed=False,
            function="text_not_contains",
            detail=f"forbidden keywords found: {found}",
        )

    @staticmethod
    def _check_excel_cell_value(output_dir: str, args: dict) -> CheckResult:
        """检查 Excel 单元格值(借鉴 SpreadsheetBench 的精确区域比对)。"""
        filename = args.get("filename", "")
        sheet_name = args.get("sheet", None)
        cell_ref = args.get("cell", "")  # 如 "A1" 或 "B3"
        expected = args.get("value", "")
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=False, function="excel_cell_value", detail=f"file not found: {filename}"
            )

        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, data_only=True)
            ws = wb[sheet_name] if sheet_name else wb.active
            actual = ws[cell_ref].value
            wb.close()

            # 归一化比对(借鉴 SpreadsheetBench 的 transform_value)
            def normalize(v):
                if v is None:
                    return ""
                if isinstance(v, (int, float)):
                    return round(float(v), 2)
                return str(v).strip()

            actual_n = normalize(actual)
            expected_n = normalize(expected)

            if actual_n == expected_n:
                return CheckResult(
                    passed=True, function="excel_cell_value",
                    detail=f"{cell_ref}={actual_n}"
                )
            return CheckResult(
                passed=False, function="excel_cell_value",
                detail=f"{cell_ref}: expected={expected_n}, actual={actual_n}"
            )
        except Exception as e:
            return CheckResult(
                passed=False, function="excel_cell_value", detail=f"error: {e}"
            )

    @staticmethod
    def _check_word_heading_exists(output_dir: str, args: dict) -> CheckResult:
        """检查 Word 文档中是否存在指定标题。"""
        filename = args.get("filename", "")
        heading_text = args.get("text", "")
        level = args.get("level", None)
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=False, function="word_heading_exists", detail=f"file not found: {filename}"
            )

        try:
            from docx import Document
            doc = Document(filepath)
            for p in doc.paragraphs:
                if p.style and p.style.name and p.style.name.startswith("Heading"):
                    if heading_text in p.text:
                        if level is None or p.style.name == f"Heading {level}":
                            return CheckResult(
                                passed=True, function="word_heading_exists",
                                detail=f"found: '{p.text}' ({p.style.name})"
                            )
            return CheckResult(
                passed=False, function="word_heading_exists",
                detail=f"heading '{heading_text}' not found"
            )
        except Exception as e:
            return CheckResult(
                passed=False, function="word_heading_exists", detail=f"error: {e}"
            )

    @staticmethod
    def _check_word_count(output_dir: str, args: dict) -> CheckResult:
        """检查 Word 文档字数。"""
        filename = args.get("filename", "")
        min_count = args.get("min", 0)
        max_count = args.get("max", 0)
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=False, function="word_word_count", detail=f"file not found: {filename}"
            )

        try:
            from docx import Document
            doc = Document(filepath)
            count = sum(len(p.text) for p in doc.paragraphs)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        count += len(cell.text)

            ok = True
            detail = f"word_count={count}"
            if min_count and count < min_count:
                ok = False
                detail += f" (below min={min_count})"
            if max_count and count > max_count:
                ok = False
                detail += f" (above max={max_count})"

            return CheckResult(passed=ok, function="word_word_count", detail=detail)
        except Exception as e:
            return CheckResult(
                passed=False, function="word_word_count", detail=f"error: {e}"
            )

    @staticmethod
    def _check_pdf_page_count(output_dir: str, args: dict) -> CheckResult:
        """检查 PDF 页数。"""
        filename = args.get("filename", "")
        expected = args.get("pages", 0)
        filepath = os.path.join(output_dir, filename)

        if not os.path.exists(filepath):
            return CheckResult(
                passed=False, function="pdf_page_count", detail=f"file not found: {filename}"
            )

        try:
            import fitz
            doc = fitz.open(filepath)
            actual = len(doc)
            doc.close()
            ok = (actual == expected) if expected else True
            return CheckResult(
                passed=ok, function="pdf_page_count",
                detail=f"pages={actual}, expected={expected}"
            )
        except Exception as e:
            return CheckResult(
                passed=False, function="pdf_page_count", detail=f"error: {e}"
            )

    @staticmethod
    def _check_file_extension(output_dir: str, args: dict) -> CheckResult:
        """检查文件扩展名。"""
        filename = args.get("filename", "")
        expected_ext = args.get("extension", "").lstrip(".").lower()
        actual_ext = os.path.splitext(filename)[1].lstrip(".").lower()
        ok = actual_ext == expected_ext
        return CheckResult(
            passed=ok, function="file_extension",
            detail=f"actual=.{actual_ext}, expected=.{expected_ext}"
        )

    def register_check(self, name: str, func: Callable) -> None:
        """注册自定义检查函数(借鉴 Unstructured 的注册表模式)。

        Args:
            name: 检查函数名
            func: 检查函数签名为 (output_dir: str, args: dict) -> CheckResult
        """
        self._check_functions[name] = func
