"""
影响溯源系统 (Impact Tracker)。

参考 Anthropic 三层防御体系的第三层(可追溯),对高风险文件操作做:
  - before(): 操作前生成快照(文件哈希 + docx 段落 + xlsx 单元格)
  - after():  操作后生成快照,与 before 对比生成 diff_summary
  - rollback(operation_id): 一键回滚到 before 快照(仅恢复内容)

快照存储: snapshot_dir/{operation_id}_{before|after}.json
可选依赖:
  - python-docx  读取 .docx 段落(缺失时 paragraphs=[])
  - openpyxl     读取 .xlsx 单元格(缺失时 cells=[])

设计原则:
  - 不修改 office/base.py,本模块独立
  - 可选依赖缺失时降级(只存哈希,不存段落/单元格)
  - 所有异常不外泄,失败时记录 warning 并返回 None / False
"""

# -*- coding: utf-8 -*-
# Copyright (C) 2026 FnixAgent. All rights reserved.
# Software Name: FnixAgent 智能工作台系统 V1.0
# This software and its source code are proprietary and confidential.
# Unauthorized copying, modification, distribution, or use is strictly prohibited.

from __future__ import annotations

import hashlib
import json
import logging
import os
import uuid
from dataclasses import dataclass, field
from datetime import datetime

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------

@dataclass
class Snapshot:
    """文件快照。

    Attributes:
        path: 文件绝对路径
        sha256: 文件内容 SHA256
        size: 文件大小(字节)
        mtime: 最后修改时间(ISO 字符串)
        paragraphs: docx 段落文本列表(若为 docx)
        cells: xlsx 单元格二维列表(若为 xlsx)
    """

    path: str
    sha256: str
    size: int
    mtime: str
    paragraphs: list[str] = field(default_factory=list)
    cells: list[list[str]] = field(default_factory=list)

@dataclass
class ImpactRecord:
    """单次操作的影响记录。

    Attributes:
        operation_id: 操作 ID(UUID4)
        timestamp: ISO 时间戳
        tool_name: 触发工具名
        target_path: 目标文件路径
        before: 操作前快照(可能为 None,如新建文件)
        after: 操作后快照(可能为 None,如删除文件)
        diff_summary: 差异摘要(新增X行/删除Y行/修改Z处)
    """

    operation_id: str
    timestamp: str
    tool_name: str
    target_path: str
    before: Snapshot | None
    after: Snapshot | None
    diff_summary: str = ""

# ---------------------------------------------------------------------------
# ImpactTracker
# ---------------------------------------------------------------------------

class ImpactTracker:
    """影响溯源器。

    用法:
        tracker = ImpactTracker(snapshot_dir="assets/snapshots")
        op_id = tracker.before("/data/report.docx", tool_name="edit_docx")
        # ... 执行编辑操作 ...
        record = tracker.after(op_id, "/data/report.docx")
        # 出错时回滚
        tracker.rollback(op_id)
    """

    # 单文件快照大小上限(50 MB,防止 OOM)
    _MAX_SNAPSHOT_SIZE = 50 * 1024 * 1024

    def __init__(self, snapshot_dir: str = "assets/snapshots") -> None:
        self._snapshot_dir = snapshot_dir
        os.makedirs(self._snapshot_dir, exist_ok=True)
        # 内存索引: operation_id → ImpactRecord(便于 after/rollback 查找)
        self._records: dict[str, ImpactRecord] = {}

    # -- 公开接口 ----------------------------------------------------------

    def before(self, path: str, tool_name: str = "unknown") -> str:
        """操作前生成快照,返回 operation_id。

        若文件不存在(如新建场景),before=None 但仍生成 operation_id。

        Args:
            path: 目标文件路径
            tool_name: 触发工具名

        Returns:
            operation_id (UUID4 hex)
        """
        op_id = uuid.uuid4().hex
        snapshot = self._take_snapshot(path)
        record = ImpactRecord(
            operation_id=op_id,
            timestamp=datetime.utcnow().isoformat(),
            tool_name=tool_name,
            target_path=os.path.realpath(path),
            before=snapshot,
            after=None,
            diff_summary="",
        )
        self._records[op_id] = record
        # 持久化 before 快照
        if snapshot is not None:
            self._save_snapshot(op_id, "before", snapshot)
        return op_id

    def after(
        self,
        operation_id: str,
        path: str,
    ) -> ImpactRecord | None:
        """操作后生成快照,计算 diff_summary 并返回 ImpactRecord。

        Args:
            operation_id: before() 返回的 ID
            path: 同一目标文件路径

        Returns:
            ImpactRecord(完整记录),若 ID 不存在返回 None
        """
        record = self._records.get(operation_id)
        if record is None:
            logger.warning("[impact] operation_id 不存在: %s", operation_id)
            return None
        after_snap = self._take_snapshot(path)
        record.after = after_snap
        record.diff_summary = self._compute_diff(record.before, after_snap)
        if after_snap is not None:
            self._save_snapshot(operation_id, "after", after_snap)
        # 更新内存记录的 timestamp 为完成时间
        record.timestamp = datetime.utcnow().isoformat()
        return record

    def rollback(self, operation_id: str) -> bool:
        """回滚到 before 快照(仅恢复内容,不恢复 mtime)。

        Args:
            operation_id: before() 返回的 ID

        Returns:
            True=成功回滚,False=无 before 快照或回滚失败
        """
        record = self._records.get(operation_id)
        if record is None:
            logger.warning("[impact] rollback: operation_id 不存在: %s", operation_id)
            return False
        if record.before is None:
            logger.warning("[impact] rollback: 无 before 快照(新建文件场景)")
            return False
        try:
            # 优先从磁盘加载 before 快照(内存可能丢失)
            before = record.before
            disk_snap = self._load_snapshot(operation_id, "before")
            if disk_snap is not None:
                before = disk_snap
            # 若 before.sha256 为空(原文件不存在但记录了占位),无法回滚
            if not before.sha256:
                return False
            # 读取磁盘快照内容(从快照 JSON 中还原 base64 内容)
            content = self._read_snapshot_content(operation_id, "before")
            if content is None:
                # 若快照没存内容(超大文件场景),无法回滚
                logger.warning("[impact] rollback: 快照未保存文件内容")
                return False
            # 写回原文件(覆盖现有内容)
            target = record.target_path or before.path
            os.makedirs(os.path.dirname(target) or ".", exist_ok=True)
            with open(target, "wb") as f:
                f.write(content)
            self._audit_rollback(operation_id, record.target_path, success=True)
            return True
        except Exception as exc:
            logger.exception("[impact] rollback 失败")
            self._audit_rollback(operation_id, record.target_path, success=False, error=str(exc))
            return False

    def list_operations(self, limit: int = 50) -> list[ImpactRecord]:
        """返回最近 N 条操作记录(按时间倒序)。"""
        records = sorted(
            self._records.values(),
            key=lambda r: r.timestamp,
            reverse=True,
        )
        return records[:limit]

    def get_record(self, operation_id: str) -> ImpactRecord | None:
        """按 ID 查询操作记录。"""
        return self._records.get(operation_id)

    # -- 内部:快照采集 ----------------------------------------------------

    def _take_snapshot(self, path: str) -> Snapshot | None:
        """采集文件快照(哈希 + 段落 + 单元格)。

        文件不存在/不可读时返回 None。
        """
        if not os.path.exists(path) or not os.path.isfile(path):
            return None
        try:
            stat = os.stat(path)
            if stat.st_size > self._MAX_SNAPSHOT_SIZE:
                logger.warning(
                    "[impact] 文件过大(%d bytes),仅记录哈希: %s",
                    stat.st_size,
                    path,
                )
                sha = self._sha256_file(path)
                return Snapshot(
                    path=os.path.realpath(path),
                    sha256=sha,
                    size=stat.st_size,
                    mtime=datetime.utcfromtimestamp(stat.st_mtime).isoformat(),
                )
            sha = self._sha256_file(path)
            paragraphs = self._read_docx_paragraphs(path)
            cells = self._read_xlsx_cells(path)
            return Snapshot(
                path=os.path.realpath(path),
                sha256=sha,
                size=stat.st_size,
                mtime=datetime.utcfromtimestamp(stat.st_mtime).isoformat(),
                paragraphs=paragraphs,
                cells=cells,
            )
        except Exception as exc:
            logger.warning("[impact] 快照采集失败 %s: %s", path, exc)
            return None

    @staticmethod
    def _sha256_file(path: str) -> str:
        """计算文件 SHA256(流式读取,避免 OOM)。"""
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(64 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _read_docx_paragraphs(path: str) -> list[str]:
        """读取 docx 段落文本(可选依赖缺失时返回 [])。"""
        if not path.lower().endswith(".docx"):
            return []
        try:
            from docx import Document  # type: ignore[import-not-found]

            doc = Document(path)
            return [p.text for p in doc.paragraphs]
        except Exception:
            return []

    @staticmethod
    def _read_xlsx_cells(path: str) -> list[list[str]]:
        """读取 xlsx 单元格(read_only 模式,缺失时返回 [])。"""
        if not path.lower().endswith(".xlsx"):
            return []
        try:
            from openpyxl import load_workbook  # type: ignore[import-not-found]

            wb = load_workbook(path, read_only=True, data_only=True)
            rows: list[list[str]] = []
            ws = wb.active
            if ws is None:
                wb.close()
                return rows
            for row in ws.iter_rows(values_only=True):
                rows.append(["" if v is None else str(v) for v in row])
            wb.close()
            return rows
        except Exception:
            return []

    # -- 内部:差异计算 ----------------------------------------------------

    def _compute_diff(
        self,
        before: Snapshot | None,
        after: Snapshot | None,
    ) -> str:
        """计算 before/after 快照差异摘要。"""
        # 文件被删除
        if before is not None and after is None:
            return "文件被删除"
        # 文件被新建
        if before is None and after is not None:
            return f"文件被新建(size={after.size})"
        # 都为 None(不应发生)
        if before is None and after is None:
            return "无变化"
        # 都存在:比较哈希
        if before.sha256 == after.sha256:
            return "无变化"
        # 内容不同:基于段落/单元格细化
        if before.paragraphs or after.paragraphs:
            return self._diff_paragraphs(before.paragraphs, after.paragraphs)
        if before.cells or after.cells:
            return self._diff_cells(before.cells, after.cells)
        # 通用:size 变化
        delta = after.size - before.size
        sign = "+" if delta >= 0 else ""
        return f"内容变化({sign}{delta} bytes)"

    @staticmethod
    def _diff_paragraphs(before: list[str], after: list[str]) -> str:
        """段落级 diff:新增/删除/修改行数。"""
        before_set = set(before)
        after_set = set(after)
        added = len(after_set - before_set)
        removed = len(before_set - after_set)
        # 修改 = before 与 after 中位置相同但内容不同的段落
        common = min(len(before), len(after))
        modified = sum(1 for i in range(common) if before[i] != after[i])
        return f"段落: 新增{added}/删除{removed}/修改{modified}"

    @staticmethod
    def _diff_cells(before: list[list[str]], after: list[list[str]]) -> str:
        """单元格级 diff:新增/删除行数 + 修改单元格数。"""
        added_rows = max(0, len(after) - len(before))
        removed_rows = max(0, len(before) - len(after))
        common_rows = min(len(before), len(after))
        modified_cells = 0
        for i in range(common_rows):
            b_row = before[i]
            a_row = after[i] if i < len(after) else []
            common_cols = min(len(b_row), len(a_row))
            for j in range(common_cols):
                if b_row[j] != a_row[j]:
                    modified_cells += 1
        return f"单元格: 新增行{added_rows}/删除行{removed_rows}/修改单元格{modified_cells}"

    # -- 内部:快照持久化 --------------------------------------------------

    def _snapshot_path(self, operation_id: str, phase: str) -> str:
        """快照文件路径: {snapshot_dir}/{op_id}_{phase}.json"""
        return os.path.join(self._snapshot_dir, f"{operation_id}_{phase}.json")

    def _save_snapshot(
        self,
        operation_id: str,
        phase: str,
        snapshot: Snapshot,
    ) -> None:
        """将快照落盘(含 base64 编码的文件内容,供 rollback 使用)。"""
        try:
            import base64

            # 仅小文件保存完整内容(供回滚),大文件只存元信息
            content_b64 = ""
            if snapshot.size <= self._MAX_SNAPSHOT_SIZE:
                try:
                    with open(snapshot.path, "rb") as f:
                        content_b64 = base64.b64encode(f.read()).decode("ascii")
                except Exception:
                    content_b64 = ""
            data = {
                "path": snapshot.path,
                "sha256": snapshot.sha256,
                "size": snapshot.size,
                "mtime": snapshot.mtime,
                "paragraphs": snapshot.paragraphs,
                "cells": snapshot.cells,
                "content_base64": content_b64,
            }
            with open(self._snapshot_path(operation_id, phase), "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False)
        except Exception:
            logger.warning("[impact] 快照落盘失败: %s_%s", operation_id, phase)

    def _load_snapshot(
        self,
        operation_id: str,
        phase: str,
    ) -> Snapshot | None:
        """从磁盘加载快照(失败返回 None)。"""
        path = self._snapshot_path(operation_id, phase)
        if not os.path.exists(path):
            return None
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            return Snapshot(
                path=data.get("path", ""),
                sha256=data.get("sha256", ""),
                size=data.get("size", 0),
                mtime=data.get("mtime", ""),
                paragraphs=data.get("paragraphs", []),
                cells=data.get("cells", []),
            )
        except Exception:
            return None

    def _read_snapshot_content(
        self,
        operation_id: str,
        phase: str,
    ) -> bytes | None:
        """从磁盘快照读取原始文件内容(base64 解码)。"""
        path = self._snapshot_path(operation_id, phase)
        if not os.path.exists(path):
            return None
        try:
            import base64

            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            content_b64 = data.get("content_base64", "")
            if not content_b64:
                return None
            return base64.b64decode(content_b64)
        except Exception:
            return None

    # -- 内部:审计 -------------------------------------------------------

    @staticmethod
    def _audit_rollback(
        operation_id: str,
        target_path: str,
        success: bool,
        error: str | None = None,
    ) -> None:
        """将回滚操作写入审计日志(失败不影响主流程)。"""
        try:
            from fnixagent.core.audit import AuditLogger

            AuditLogger().log(
                action="impact.rollback",
                detail={
                    "operation_id": operation_id,
                    "target_path": target_path,
                    "success": success,
                    "error": error,
                },
            )
        except Exception:
            pass
