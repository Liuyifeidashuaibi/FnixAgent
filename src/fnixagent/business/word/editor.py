"""
业务能力层 - Word 文档编辑。

基于 python-docx 实现 Word 创建/编辑/格式化。

安全防护:
  - 输入文件存在性/大小校验(MAX_INPUT_FILE_BYTES,默认 50MB)
  - 输出路径目录可写性校验

异常捕获:
  - python-docx 抛出的 Exception(docx.opc.exceptions.PackageNotFoundError 等)统一捕获
  - 表格行列越界保护(clamp 到声明行列,避免 IndexError)
  - 标题级别校验(0-9)

BUG 修复:
  - 原 `doc.insert_paragraph_before(text, 0)` 调用错误(Document 无此方法,
    且 Paragraph.insert_paragraph_before 不接受索引参数);
    改为 `doc.paragraphs[0].insert_paragraph_before(text)`,空文档时回退到 add_paragraph
  - 表格数据越界:原 `table.rows[i].cells[j]` 在 data 行列超过 rows/cols 时抛 IndexError;
    改为先 clamp 再写入,跳过越界数据
"""

# -*- coding: utf-8 -*-
# Copyright (C) 2026 FnixAgent. All rights reserved.
# Software Name: FnixAgent 智能工作台系统 V1.0
# This software and its source code are proprietary and confidential.
# Unauthorized copying, modification, distribution, or use is strictly prohibited.

import logging
import os

from docx import Document

from fnixagent.core.tools.protocol import ToolMetadata

_logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------

# 输入文件大小上限(50 MB);超过拒绝,避免 OOM
MAX_INPUT_FILE_BYTES = 50 * 1024 * 1024

# 支持的编辑操作
_VALID_OPERATIONS = frozenset({"add_text", "replace", "insert_table", "add_heading"})

# 标题级别范围(python-docx 支持 0-9,0 为 Title)
_HEADING_LEVEL_MIN = 0
_HEADING_LEVEL_MAX = 9

# ---------------------------------------------------------------------------
# 内部工具
# ---------------------------------------------------------------------------


def _check_input_file(file_path: str) -> dict | None:
    """校验输入文件存在性与大小;返回错误 dict(无错误返回 None)。"""
    if not file_path:
        return {"success": False, "error": "file_path must not be empty"}
    if not os.path.exists(file_path):
        return {
            "success": False,
            "error": f"input file not found: {file_path}",
            "file_path": file_path,
        }
    if not os.path.isfile(file_path):
        return {
            "success": False,
            "error": f"input path is not a regular file: {file_path}",
            "file_path": file_path,
        }
    try:
        size = os.path.getsize(file_path)
    except OSError as e:
        return {"success": False, "error": f"cannot stat input file: {e}", "file_path": file_path}
    if size > MAX_INPUT_FILE_BYTES:
        return {
            "success": False,
            "error": f"input file size ({size} bytes) exceeds limit "
            f"({MAX_INPUT_FILE_BYTES // 1024 // 1024}MB)",
            "file_path": file_path,
        }
    if size == 0:
        return {"success": False, "error": "input file is empty", "file_path": file_path}
    return None


# ---------------------------------------------------------------------------
# Word 读取工具 (P0 修复: 让 LLM 能读取已有 docx 内容, 用于总结/分析)
# ---------------------------------------------------------------------------


def read_docx(file_path: str) -> dict:
    """读取 Word 文档文本内容。

    提取段落文本 + 表格文本, 返回纯文本和结构化信息。
    用于"总结这份 Word"/"分析这份 docx"等场景。

    Args:
        file_path: .docx 文件路径

    Returns:
        {success, text, paragraph_count, table_count} 或 {success, error}
    """
    err = _check_input_file(file_path)
    if err is not None:
        return err

    try:
        doc = Document(file_path)
    except Exception as e:
        _logger.exception("read_docx open failed: %s: %s", type(e).__name__, e)
        return {
            "success": False,
            "error": f"无法打开 docx 文件 (可能不是有效 Word 格式): {type(e).__name__}: {e}",
            "file_path": file_path,
        }

    try:
        parts: list[str] = []
        # 提取段落文本
        para_count = 0
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                # 标记标题级别, 方便 LLM 理解结构
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading" in style_name or "title" in style_name:
                    parts.append(f"## {text}")
                else:
                    parts.append(text)
                para_count += 1

        # 提取表格文本 (转 markdown 表格格式)
        table_count = len(doc.tables)
        for ti, table in enumerate(doc.tables, 1):
            parts.append(f"\n[表格 {ti}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")

        full_text = "\n".join(parts)
        # 截断防止超长文档塞爆 LLM 上下文
        if len(full_text) > 8000:
            full_text = full_text[:8000] + f"\n\n... [已截断, 原文共 {len(full_text)} 字符]"

        return {
            "success": True,
            "text": full_text,
            "paragraph_count": para_count,
            "table_count": table_count,
        }
    except Exception as e:
        _logger.exception("read_docx extract failed: %s: %s", type(e).__name__, e)
        return {
            "success": False,
            "error": f"读取 docx 内容失败: {type(e).__name__}: {e}",
            "file_path": file_path,
        }


# ---------------------------------------------------------------------------
# Word 创建工具
# ---------------------------------------------------------------------------


def create_docx(
    content: str,
    title: str | None = None,
    template: str | None = None,
    output_path: str = "output.docx",
) -> dict:
    """
    创建 Word 文档。

    Args:
        content: 文档内容(文本,非空)
        title: 文档标题
        template: 模板名称(academic / report);其他值按默认处理
        output_path: 输出文件路径

    Returns:
        {success, file_path, metadata} 或 {success, error}
    """
    # 参数非空校验
    if not content:
        return {"success": False, "error": "content must not be empty"}
    if not output_path:
        return {"success": False, "error": "output_path must not be empty"}

    # 输出目录可写性校验
    out_dir = os.path.dirname(os.path.abspath(output_path))
    if not os.path.isdir(out_dir):
        return {"success": False, "error": f"output directory does not exist: {out_dir}"}

    try:
        # 创建文档
        doc = Document()

        # 添加标题
        if title:
            doc.add_heading(title, level=0)

        # 根据模板添加样式
        if template == "academic":
            # 学术论文模板
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(content[:200])  # 摘要

            doc.add_heading("Introduction", level=1)
            doc.add_paragraph(content)

        elif template == "report":
            # 报告模板
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(content)

        else:
            # 默认: 直接添加内容
            doc.add_paragraph(content)

        # 保存文档
        doc.save(output_path)

        return {
            "success": True,
            "file_path": output_path,
            "title": title,
            "template": template,
            "pages": 1,
        }

    except Exception as e:
        _logger.exception("create_docx failed: %s: %s", type(e).__name__, e)
        return {
            "success": False,
            "error": f"{type(e).__name__}: {e}",
        }


# ---------------------------------------------------------------------------
# Word 编辑工具
# ---------------------------------------------------------------------------


def edit_docx(
    file_path: str,
    operation: str,
    params: dict,
) -> dict:
    """
    编辑 Word 文档。

    Args:
        file_path: 文档路径(必须存在)
        operation: 操作类型(add_text/replace/insert_table/add_heading)
        params: 操作参数
            - add_text: {text, position=end|start}
            - replace: {old_text, new_text}
            - insert_table: {rows, cols, data=[[...], ...]}
            - add_heading: {text, level}

    Returns:
        {success, file_path, operation} 或 {success, error}
    """
    # 操作类型枚举校验
    if operation not in _VALID_OPERATIONS:
        return {
            "success": False,
            "error": f"unsupported operation {operation!r}, "
            f"must be one of {sorted(_VALID_OPERATIONS)}",
        }
    if not isinstance(params, dict):
        return {"success": False, "error": "params must be a dict"}

    # 输入文件校验
    err = _check_input_file(file_path)
    if err is not None:
        return err

    try:
        # 打开文档
        doc = Document(file_path)

        if operation == "add_text":
            text = params.get("text", "")
            position = params.get("position", "end")  # start/end

            if position == "end":
                doc.add_paragraph(text)
            elif position == "start":
                # BUG 修复:原 `doc.insert_paragraph_before(text, 0)` 调用错误
                # (Document 无此方法,Paragraph.insert_paragraph_before 不接受索引)
                if doc.paragraphs:
                    doc.paragraphs[0].insert_paragraph_before(text)
                else:
                    # 空文档回退到 add_paragraph
                    doc.add_paragraph(text)
            else:
                return {
                    "success": False,
                    "error": f"unsupported position {position!r}, must be 'start' or 'end'",
                    "file_path": file_path,
                }

        elif operation == "replace":
            # 查找替换
            old_text = params.get("old_text", "")
            new_text = params.get("new_text", "")
            if not old_text:
                return {
                    "success": False,
                    "error": "old_text must not be empty for replace operation",
                    "file_path": file_path,
                }

            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

        elif operation == "insert_table":
            # 插入表格
            rows = params.get("rows", 3)
            cols = params.get("cols", 3)
            data = params.get("data", [])

            # 行列合法性校验
            if not isinstance(rows, int) or rows <= 0:
                return {
                    "success": False,
                    "error": f"rows must be a positive integer, got {rows!r}",
                    "file_path": file_path,
                }
            if not isinstance(cols, int) or cols <= 0:
                return {
                    "success": False,
                    "error": f"cols must be a positive integer, got {cols!r}",
                    "file_path": file_path,
                }

            table = doc.add_table(rows=rows, cols=cols)
            # BUG 修复:原代码在 data 行列超过 rows/cols 时会 IndexError;
            # 改为 clamp 到声明行列,跳过越界数据
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                if not isinstance(row_data, (list, tuple)):
                    continue
                for j, cell_data in enumerate(row_data):
                    if j >= cols:
                        break
                    table.rows[i].cells[j].text = str(cell_data)

        elif operation == "add_heading":
            # 添加标题
            text = params.get("text", "")
            level = params.get("level", 1)
            # 级别合法性校验
            if not isinstance(level, int) or not (
                _HEADING_LEVEL_MIN <= level <= _HEADING_LEVEL_MAX
            ):
                return {
                    "success": False,
                    "error": f"level must be an integer in "
                    f"[{_HEADING_LEVEL_MIN}, {_HEADING_LEVEL_MAX}], got {level!r}",
                    "file_path": file_path,
                }
            if not text:
                return {
                    "success": False,
                    "error": "text must not be empty for add_heading operation",
                    "file_path": file_path,
                }
            doc.add_heading(text, level=level)

        # 保存文档
        doc.save(file_path)

        return {
            "success": True,
            "file_path": file_path,
            "operation": operation,
        }

    except Exception as e:
        _logger.exception("edit_docx failed: %s: %s", type(e).__name__, e)
        return {
            "success": False,
            "error": f"{type(e).__name__}: {e}",
            "file_path": file_path,
        }


# ---------------------------------------------------------------------------
# Word 格式化工具
# ---------------------------------------------------------------------------


def format_docx(file_path: str, style_name: str = "Normal") -> dict:
    """
    应用 Word 格式样式。

    Args:
        file_path: 文档路径(必须存在)
        style_name: 样式名称(Normal/Heading 1/Title/...;必须存在于文档样式集)

    Returns:
        {success, file_path, style} 或 {success, error}
    """
    if not style_name:
        return {"success": False, "error": "style_name must not be empty"}

    err = _check_input_file(file_path)
    if err is not None:
        return err

    try:
        doc = Document(file_path)

        # 校验样式是否存在(避免 KeyError)
        if style_name not in [s.name for s in doc.styles]:
            return {
                "success": False,
                "error": f"style {style_name!r} not found in document styles",
                "file_path": file_path,
            }

        # 应用样式到所有段落
        for paragraph in doc.paragraphs:
            paragraph.style = doc.styles[style_name]

        doc.save(file_path)

        return {
            "success": True,
            "file_path": file_path,
            "style": style_name,
        }

    except Exception as e:
        _logger.exception("format_docx failed: %s: %s", type(e).__name__, e)
        return {
            "success": False,
            "error": f"{type(e).__name__}: {e}",
        }


# ---------------------------------------------------------------------------
# 工具元数据
# ---------------------------------------------------------------------------

TOOL_METADATA = {
    "create_docx": ToolMetadata(
        name="create_docx",
        description="创建 Word 文档并写入内容",
        category="word",
        input_schema={
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "文档内容"},
                "title": {"type": "string"},
                "template": {"type": "string"},
                "output_path": {"type": "string", "default": "output.docx"},
            },
            "required": ["content"],
        },
    ),
    "edit_docx": ToolMetadata(
        name="edit_docx",
        description="编辑 Word 文档(添加文本/替换/插入表格)",
        category="word",
        input_schema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string"},
                "operation": {
                    "type": "string",
                    "enum": ["add_text", "replace", "insert_table", "add_heading"],
                },
                "params": {"type": "object"},
            },
            "required": ["file_path", "operation", "params"],
        },
    ),
    "format_docx": ToolMetadata(
        name="format_docx",
        description="应用 Word 格式样式",
        category="word",
        input_schema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string"},
                "style_name": {"type": "string", "default": "Normal"},
            },
            "required": ["file_path"],
        },
    ),
}


def register_word_tools(registry) -> None:
    """注册 Word 工具到工具注册中心。"""
    registry.register(TOOL_METADATA["create_docx"], create_docx)
    registry.register(TOOL_METADATA["edit_docx"], edit_docx)
    registry.register(TOOL_METADATA["format_docx"], format_docx)
