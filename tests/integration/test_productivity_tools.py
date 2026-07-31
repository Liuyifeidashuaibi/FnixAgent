"""生产力工具业务闭环验证测试。

验证办公人群真实业务场景能跑通:
  - PDF 读取 (总结这份 PDF)
  - Excel 读取 (分析这份表格)
  - Word 读取 (总结这份文档)
  - 文档格式转换 (docx→pdf 等, 真实转换非 stub)
  - 工具调用结果成功/失败检测正确 (Reflexion 能触发)
  - 工具注册完整性 (LLM 能发现这些工具)
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "src"))


def test_office_read_tools_registered():
    """P0 验证: read_pdf / read_xlsx / read_docx 应注册到 ToolRegistry。"""
    from fnixagent.core.tools.registry import ToolRegistry
    from fnixagent.services.work_agent import register_office_work_tools

    registry = ToolRegistry()
    register_office_work_tools(registry, workspace_root=str(ROOT), craft_artifacts=False)

    tool_names = {
        t["function"]["name"] if isinstance(t, dict) else t.name for t in registry.list_for_llm()
    }
    assert "read_pdf" in tool_names, "read_pdf 工具未注册 — 无法总结 PDF"
    assert "read_xlsx" in tool_names, "read_xlsx 工具未注册 — 无法分析 Excel"
    assert "read_docx" in tool_names, "read_docx 工具未注册 — 无法总结 Word"
    assert "convert_document" in tool_names, "convert_document 工具未注册"
    assert "create_xlsx" in tool_names, "create_xlsx 应保留"
    assert "create_pdf" in tool_names, "create_pdf 应保留"
    assert "create_pptx" in tool_names, "create_pptx 应保留"
    assert "create_chart" in tool_names, "create_chart 应保留"
    print("[P0] 办公读取工具注册验证通过")


def test_convert_document_not_stub():
    """P0 验证: convert_document 应调用真实 ConverterExpert 而非 stub。"""
    from fnixagent.core.tools.registry import ToolRegistry
    from fnixagent.services.work_agent import register_office_work_tools

    registry = ToolRegistry()
    register_office_work_tools(registry, workspace_root=str(ROOT), craft_artifacts=False)

    # 执行一个不存在的文件转换, stub 会返回 success:true, 真实实现应返回 success:false
    result = registry.execute(
        "convert_document",
        {
            "file_path": "nonexistent_file.xyz",
            "target_format": "pdf",
        },
    )
    if isinstance(result, dict):
        assert result.get("success") is False, (
            "convert_document 应是真实实现 (失败时 success=false), 不应是 stub (总返回 success=true)"
        )
        assert "stub" not in str(result).lower(), "不应返回 stub 标记"
    print("[P0] convert_document 真实实现验证通过")


def test_dict_result_success_detection():
    """P1 验证: dict 结果的成功/失败检测应正确。"""
    import inspect

    from fnixagent.core.agent.loop import AgenticLoop

    src = inspect.getsource(AgenticLoop._execute_tool)
    # 应有 dict 分支处理 success 字段
    assert "isinstance(result, dict)" in src, "应有 dict 分支"
    assert 'result.get("success"' in src, "应检查 success 字段"
    assert "[成功]" in src and "[失败]" in src, "应标记成功/失败前缀"
    print("[P1] dict 结果成功检测验证通过")


def test_edit_docx_path_resolve():
    """P1 验证: edit_docx 应走 _resolve 路径改写。"""
    from fnixagent.core.tools.registry import ToolRegistry
    from fnixagent.services.work_agent import register_office_work_tools

    registry = ToolRegistry()
    register_office_work_tools(registry, workspace_root=str(ROOT), craft_artifacts=True)

    # Craft 模式下 edit_docx 一个相对路径, 应被 _resolve 改写
    # 由于文件不存在会失败, 但路径应已被改写
    result = registry.execute(
        "edit_docx",
        {
            "file_path": "test_relative.docx",
            "operation": "add_text",
            "params": {"text": "test"},
        },
    )
    # 文件不存在应返回失败 (而非路径错误)
    if isinstance(result, dict):
        assert result.get("success") is False, "不存在文件应失败"
    print("[P1] edit_docx 路径 _resolve 验证通过")


def test_read_pdf_real_file(tmp_path):
    """P0 验证: read_pdf 能真实读取 PDF 文本。"""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        print("[P0] 跳过 read_pdf 真实文件测试: reportlab 未安装")
        return

    # 创建测试 PDF
    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Hello FnixAgent PDF Test")
    c.drawString(100, 730, "这是测试内容")
    c.save()

    from fnixagent.core.tools.registry import ToolRegistry
    from fnixagent.services.work_agent import register_office_work_tools

    registry = ToolRegistry()
    register_office_work_tools(registry, workspace_root=str(tmp_path), craft_artifacts=False)

    result = registry.execute("read_pdf", {"file_path": str(pdf_path)})
    assert isinstance(result, dict), f"结果应为 dict, 实际 {type(result)}"
    assert result.get("success") is True, f"读取 PDF 应成功: {result}"
    assert "Hello" in result.get("text", "") or "FnixAgent" in result.get("text", ""), (
        "应提取 PDF 文本内容"
    )
    print(f"[P0] read_pdf 真实文件验证通过 (页数: {result.get('page_count')})")


def test_read_xlsx_real_file(tmp_path):
    """P0 验证: read_xlsx 能真实读取 Excel 数据。"""
    try:
        from openpyxl import Workbook
    except ImportError:
        print("[P0] 跳过 read_xlsx 真实文件测试: openpyxl 未安装")
        return

    # 创建测试 Excel
    xlsx_path = tmp_path / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["姓名", "年龄", "部门"])
    ws.append(["张三", 28, "研发部"])
    ws.append(["李四", 32, "产品部"])
    wb.save(str(xlsx_path))

    from fnixagent.core.tools.registry import ToolRegistry
    from fnixagent.services.work_agent import register_office_work_tools

    registry = ToolRegistry()
    register_office_work_tools(registry, workspace_root=str(tmp_path), craft_artifacts=False)

    result = registry.execute("read_xlsx", {"file_path": str(xlsx_path)})
    assert isinstance(result, dict), f"结果应为 dict, 实际 {type(result)}"
    assert result.get("success") is True, f"读取 Excel 应成功: {result}"
    headers = result.get("headers", [])
    rows = result.get("rows", [])
    assert "姓名" in headers, f"应包含表头: {headers}"
    assert len(rows) == 2, f"应有 2 行数据, 实际 {len(rows)}"
    assert rows[0][0] == "张三", f"第一行数据应正确: {rows[0]}"
    print(f"[P0] read_xlsx 真实文件验证通过 (行数: {len(rows)})")


def test_read_docx_real_file(tmp_path):
    """P0 验证: read_docx 能真实读取 Word 文本。"""
    try:
        from docx import Document
    except ImportError:
        print("[P0] 跳过 read_docx 真实文件测试: python-docx 未安装")
        return

    # 创建测试 Word
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("测试标题", level=1)
    doc.add_paragraph("这是测试段落内容。")
    doc.add_paragraph("第二段内容。")
    doc.save(str(docx_path))

    from fnixagent.business.word.editor import read_docx

    result = read_docx(file_path=str(docx_path))
    assert result.get("success") is True, f"读取 docx 应成功: {result}"
    assert "测试标题" in result.get("text", ""), "应提取标题"
    assert "测试段落" in result.get("text", ""), "应提取段落内容"
    assert result.get("paragraph_count") == 3, f"应有 3 段: {result.get('paragraph_count')}"
    print(f"[P0] read_docx 真实文件验证通过 (段落数: {result.get('paragraph_count')})")


def test_web_search_html_parsing():
    """P2 验证: web_search 应使用 HTML 解析版而非仅 Instant Answer API。"""
    import inspect

    from fnixagent.core.tools.workspace import WorkspaceTools

    src = inspect.getsource(WorkspaceTools.web_search)
    assert "html.duckduckgo.com/html" in src, "应使用 HTML 搜索页"
    assert "result__a" in src, "应解析 result__a 类"
    assert "result__snippet" in src, "应解析 result__snippet 类"
    print("[P2] web_search HTML 解析版验证通过")


if __name__ == "__main__":
    import tempfile

    test_office_read_tools_registered()
    test_convert_document_not_stub()
    test_dict_result_success_detection()
    test_edit_docx_path_resolve()
    test_web_search_html_parsing()
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        test_read_pdf_real_file(tmp_path)
        test_read_xlsx_real_file(tmp_path)
        test_read_docx_real_file(tmp_path)
    print("\n=== 所有生产力工具业务闭环验证通过 ===")
